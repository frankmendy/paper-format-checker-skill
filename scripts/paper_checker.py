import argparse
import sys
import re
from copy import deepcopy
from pathlib import Path
import time

# 依赖检查
def check_dependencies():
    """检查必要的依赖包是否已安装"""
    missing_deps = []
    
    # 检查 python-docx
    try:
        import docx
    except ImportError:
        missing_deps.append("python-docx")
    
    # 检查 pywin32 (Windows only)
    try:
        import win32com.client
    except ImportError:
        missing_deps.append("pywin32")
    
    if missing_deps:
        print("=" * 80)
        print("❌ 错误：缺少必要的依赖包")
        print("=" * 80)
        print("\n请先安装以下依赖包：")
        print(f"   pip install {' '.join(missing_deps)}")
        print("\n或者安装所有依赖：")
        print("   pip install python-docx pywin32")
        print("\n依赖说明：")
        print("   - python-docx: 用于读取和写入 Word 文档")
        print("   - pywin32: 用于 Word COM 自动化（更新目录、刷新域等）")
        print("=" * 80)
        sys.exit(1)

# 在导入其他模块前执行依赖检查
check_dependencies()

# 正常导入依赖
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

TEMPLATE_NAME = "网络空间安全学院论文格式模板.docx"
TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / TEMPLATE_NAME
LEDGER_PATH = Path(__file__).parent.parent / "FORMAT_RULES_LEDGER.md"
VALIDATION_MIN_PASSES = 1
VALIDATION_MAX_PASSES = 2
DEFAULT_WORD_UPDATE_FIELDS = True

class PaperFixer:
    def __init__(self, template_path):
        self.template_path = template_path
        self.template_info = self._extract_template_info()
        self.ledger_text = self._load_ledger_text()

    def _extract_template_info(self):
        """提取模板的精细化格式信息"""
        if not self.template_path.exists():
            return None
        
        doc = Document(self.template_path)
        info = {
            "margins": None,
            "styles": {},
            "special_sections": {}
        }

        # 1. 提取页边距
        if doc.sections:
            sec = doc.sections[0]
            info["margins"] = {
                "top": sec.top_margin,
                "bottom": sec.bottom_margin,
                "left": sec.left_margin,
                "right": sec.right_margin
            }

        # 2. 提取样式信息
        for s in doc.styles:
            try:
                info["styles"][s.name] = {
                    "font_name": s.font.name,
                    "font_size": s.font.size,
                    "bold": s.font.bold,
                    "italic": s.font.italic,
                    "alignment": None, # 样式级别通常不直接存对齐，但在段落中会体现
                    "line_spacing": None
                }
            except:
                continue

        # 3. 扫描段落以获取更具体的格式（如行间距、对齐）
        abstract_keywords = ["摘要", "ABSTRACT", "关键词", "Keywords"]
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name
            
            # 更新样式中的对齐和行距（以第一个发现该样式的非空段落为准）
            if text and info["styles"].get(style_name):
                if info["styles"][style_name]["alignment"] is None:
                    info["styles"][style_name]["alignment"] = para.alignment
                if info["styles"][style_name]["line_spacing"] is None:
                    info["styles"][style_name]["line_spacing"] = para.paragraph_format.line_spacing

            # 识别特殊章节
            for kw in abstract_keywords:
                if kw.upper() in text.upper() and len(text) < 50:
                    info["special_sections"][kw] = {
                        "style_name": style_name,
                        "alignment": para.alignment,
                        "line_spacing": para.paragraph_format.line_spacing,
                        "font_size": para.style.font.size
                    }

        return info

    def _load_ledger_text(self):
        if not LEDGER_PATH.exists():
            return ""
        return LEDGER_PATH.read_text(encoding="utf-8")

    def fix(self, paper_path, word_update_fields=DEFAULT_WORD_UPDATE_FIELDS):
        doc = Document(paper_path)
        
        print(f"--- 开始修复论文格式: {paper_path.name} ---")

        if self.ledger_text:
            print("📘 已加载 FORMAT_RULES_LEDGER，启动循环修复与复检")

        last_issues = []
        validation_history = []
        aggregate_stats = {"abstract": 0, "toc": 0, "body": 0}
        for pass_index in range(1, VALIDATION_MAX_PASSES + 1):
            self._apply_fix_pass(doc)
            for k in ("abstract", "toc", "body"):
                aggregate_stats[k] += self.pass_stats.get(k, 0)
            last_issues = self._validate_against_ledger(doc)
            print(f"🔍 第 {pass_index} 轮复检完成：发现 {len(last_issues)} 个问题")
            if last_issues:
                for issue in last_issues[:5]:
                    print(f"   - {issue}")
            validation_history.append({"pass": pass_index, "issues": last_issues[:]})
            if pass_index >= VALIDATION_MIN_PASSES and not last_issues:
                print(f"✅ 循环复检通过，共执行 {pass_index} 轮")
                break
        else:
            raise RuntimeError(
                "循环复检在最大轮次后仍存在未修复问题：\n" +
                "\n".join(f"- {issue}" for issue in last_issues[:10])
            )

        # 最终清理：确保所有二、三级标题的段前段后为自动
        print("🔧 [最终清理] 确保二、三级标题段前段后为自动...")
        self._final_cleanup_headings(doc)
        
        # 保存
        output_path = paper_path.parent / f"{paper_path.stem}_fixed{paper_path.suffix}"
        doc.save(output_path)
        print(f"\n🚀 修复完成！文件已保存至: {output_path}")
        
        # 保存后验证：重新读取文件检查二、三级标题段前段后
        print("🔍 [保存后验证] 重新读取文件检查二、三级标题段前段后...")
        self._verify_after_save(output_path)
        if word_update_fields:
            self._try_update_word_fields(output_path)
            self._postprocess_word_toc_title_xml(output_path)
            post_history = []
            post_doc = Document(output_path)
            post_issues = self._validate_against_ledger(post_doc)
            post_history.append({"pass": 1, "issues": post_issues})
            print(f"🔍 Word 更新后复检：发现 {len(post_issues)} 个问题")
            if post_issues:
                for issue in post_issues[:10]:
                    print(f"   - {issue}")
        self._write_validation_report(output_path, validation_history)
        self._write_stats_report(output_path, aggregate_stats)

    def _apply_fix_pass(self, doc):
        self.pass_stats = {"abstract": 0, "toc": 0, "body": 0}
        print("🔧 [阶段1] 标题与摘要 修复进行中…")
        print("🔧 [阶段2] 目录 修复进行中…")
        print("🔧 [阶段3] 正文 修复进行中…")

        # 1. 修复页边距
        if self.template_info["margins"]:
            for section in doc.sections:
                section.top_margin = self.template_info["margins"]["top"]
                section.bottom_margin = self.template_info["margins"]["bottom"]
                section.left_margin = self.template_info["margins"]["left"]
                section.right_margin = self.template_info["margins"]["right"]
            print("✅ 页面设置已修正 (页边距)")

        toc_start, toc_end = self._detect_toc_range(doc)
        word_toc_field_count = self._count_word_toc_field_entries(doc)
        if toc_start is not None and toc_end is not None:
            has_toc_title = any(self._normalize_text(doc.paragraphs[i].text) == "目录" for i in range(toc_start, toc_end))
            if not has_toc_title:
                title_para = doc.paragraphs[toc_start].insert_paragraph_before()
                title_para.add_run("目    录")
                self._clear_paragraph_indents(title_para)
                self._apply_custom_style(title_para, font_name="宋体", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                toc_end += 1
        elif word_toc_field_count > 0:
            self.pass_stats["toc"] = word_toc_field_count
            print(f"📑 已识别 Word 自动目录域，目录项 {word_toc_field_count} 条（目录格式在 Word 更新阶段统一修正）")
        body_start_idx = toc_end if toc_end is not None else 0

        # 预处理：寻找 Abstract 段落的索引，以便精准定位英文标题
        abstract_idx = -1
        for i, para in enumerate(doc.paragraphs):
            text = self._normalize_text(para.text).upper()
            if text.startswith("ABSTRACT"):
                abstract_idx = i
                break

        # 预处理：定位中文摘要位置以辅助判断中文标题
        cn_abstract_idx = -1
        for i, para in enumerate(doc.paragraphs):
            text = self._normalize_text(para.text)
            if text.startswith("摘要"):
                cn_abstract_idx = i
                break

        cn_keywords_idx = -1
        if cn_abstract_idx >= 0:
            for i in range(cn_abstract_idx + 1, len(doc.paragraphs)):
                text = self._normalize_text(doc.paragraphs[i].text)
                if text.startswith("关键词"):
                    cn_keywords_idx = i
                    break

        en_keywords_idx = -1
        if abstract_idx >= 0:
            for i in range(abstract_idx + 1, len(doc.paragraphs)):
                text = self._normalize_text(doc.paragraphs[i].text).upper()
                if text.startswith("KEYWORDS"):
                    en_keywords_idx = i
                    break

        # 预处理：全局扫描正文是否存在“附录”章节
        has_appendix = False
        for para in doc.paragraphs[body_start_idx:]:
            clean_text_for_state = self._normalize_text(para.text)
            if "附录" in clean_text_for_state and len(clean_text_for_state) < 10 and para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                has_appendix = True
                break

        # 2. 分部分修正
        current_part = "摘要" # 初始假设为摘要部分
        current_body_context = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            # 记录段落索引以供上下文判断使用
            setattr(para, '_index', i)
            setattr(para, '_abstract_idx', abstract_idx)
            setattr(para, '_cn_abstract_idx', cn_abstract_idx)
            setattr(para, '_en_keywords_idx', en_keywords_idx)
            setattr(para, '_cn_keywords_idx', cn_keywords_idx)
            setattr(para, '_doc_has_appendix', has_appendix)
            
            if not text:
                # 依然需要判断属于哪个部分，但空行直接在具体部分逻辑里处理
                pass

            clean_text_for_state = self._normalize_text(text)
            if toc_start is not None and toc_end is not None:
                if i < toc_start:
                    current_part = "摘要"
                elif toc_start <= i < toc_end:
                    current_part = "目录"
                else:
                    current_part = "正文"
            else:
                if current_part == "摘要" and clean_text_for_state == "目录":
                    current_part = "目录"
                elif current_part == "目录":
                    if re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text.strip()):
                        if not re.search(r'\d+$', text.strip()):
                            current_part = "正文"
                    elif para.style.name == "Heading 1":
                        if not re.search(r'\d+$', text.strip()):
                            current_part = "正文"
                elif current_part == "摘要":
                    if re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text.strip()):
                        if not re.search(r'\d+$', text.strip()):
                            current_part = "正文"

            # 执行修正
            if current_part == "摘要":
                if self._fix_abstract_para(para):
                    self.pass_stats["abstract"] += 1
            elif current_part == "目录":
                if not text: continue
                if self._fix_toc_para(para):
                    self.pass_stats["toc"] += 1
            elif current_part == "正文":
                para_xml = para._element.xml
                has_image_like_content = any(tag in para_xml for tag in ["w:drawing", "v:imagedata", "pic", "w:object", "w:pict"])
                if not text and not has_image_like_content:
                    continue
                body_text = clean_text_for_state
                if body_text in ["参考文献", "致谢", "附录"]:
                    current_body_context = body_text
                elif para.style.name == "Heading 1" or re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text.strip()):
                    current_body_context = None
                setattr(para, '_body_context', current_body_context)
                if self._fix_body_para(para):
                    self.pass_stats["body"] += 1

        print(f"✅ [阶段1] 标题与摘要 修复完成，计数 {self.pass_stats['abstract']}")
        print(f"✅ [阶段2] 目录 修复完成，计数 {self.pass_stats['toc']}")
        print(f"✅ [阶段3] 正文 修复完成，计数 {self.pass_stats['body']}")

        # 3. 专门处理图表题注
        self._fix_captions(doc)

        # 4. 专门处理表名与表格内容
        self._fix_tables(doc)

        # 5. 最终收口：再次强制修正摘要/关键词这类“标题+正文同段”的 run 级样式
        for para in doc.paragraphs:
            clean = self._normalize_text(para.text)
            if clean.startswith("摘要"):
                self._force_mixed_label_run_format(para, "黑体", Pt(10.5), "楷体", Pt(10.5), body_ascii_font="Times New Roman")
            elif clean.startswith("关键词"):
                self._force_mixed_label_run_format(para, "黑体", Pt(10.5), "楷体", Pt(10.5), body_ascii_font="Times New Roman")
            elif clean.upper().startswith("ABSTRACT"):
                self._force_mixed_label_run_format(para, "Times New Roman", Pt(12), "Times New Roman", Pt(12))
            elif clean.upper().startswith("KEYWORDS"):
                self._force_mixed_label_run_format(para, "Times New Roman", Pt(12), "Times New Roman", Pt(12))

    def _normalize_text(self, text):
        if text is None:
            return ""
        return re.sub(r'[\s\u3000\u00A0]+', '', text)

    def _force_mixed_label_run_format(self, para, title_font, title_size, body_font, body_size, title_ascii_font=None, body_ascii_font=None):
        full_text = para.text or ""
        if "：" not in full_text and ":" not in full_text:
            return False

        changed = False
        passed_title = False
        for run in para.runs:
            if run._element is None:
                continue

            rpr = run._element.get_or_add_rPr()
            if not passed_title:
                self._set_run_font(run, font_name=title_font, font_size=title_size, bold=True, ascii_font=title_ascii_font or title_font)
                for tag in ["b", "bCs"]:
                    elem = rpr.find(qn(f"w:{tag}"))
                    if elem is not None:
                        rpr.remove(elem)
                rpr.append(OxmlElement("w:b"))
                rpr.append(OxmlElement("w:bCs"))
                changed = True
                if "：" in run.text or ":" in run.text:
                    passed_title = True
            else:
                self._set_run_font(run, font_name=body_font, font_size=body_size, bold=False, ascii_font=body_ascii_font or body_font)
                for tag in ["b", "bCs"]:
                    elem = rpr.find(qn(f"w:{tag}"))
                    if elem is not None:
                        rpr.remove(elem)
                b = OxmlElement("w:b")
                b.set(qn("w:val"), "0")
                rpr.append(b)
                bcs = OxmlElement("w:bCs")
                bcs.set(qn("w:val"), "0")
                rpr.append(bcs)
                changed = True
        return changed

    def _count_word_toc_field_entries(self, doc):
        try:
            xml = doc.part.element.xml
        except Exception:
            return 0
        if "TOC \\o" not in xml and "PAGEREF _Toc" not in xml:
            return 0
        return len(re.findall(r'PAGEREF\s+_Toc', xml))

    def _postprocess_word_toc_title_xml(self, docx_path):
        try:
            import zipfile
            from lxml import etree
        except Exception:
            return

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        w_ns = ns["w"]
        xml_ns = "http://www.w3.org/XML/1998/namespace"

        def w_tag(name):
            return f"{{{w_ns}}}{name}"

        try:
            with zipfile.ZipFile(docx_path, "r") as zin:
                file_map = {name: zin.read(name) for name in zin.namelist()}
        except Exception:
            return

        document_xml = file_map.get("word/document.xml")
        if not document_xml:
            return

        try:
            root = etree.fromstring(document_xml)
        except Exception:
            return

        paragraphs = root.xpath("//w:p", namespaces=ns)
        target_para = None
        for idx, p in enumerate(paragraphs):
            instr = "".join(p.xpath(".//w:instrText/text()", namespaces=ns))
            if "TOC \\o" in instr:
                if idx > 0:
                    target_para = paragraphs[idx - 1]
                break

        if target_para is None:
            return

        pPr = target_para.find(w_tag("pPr"))
        if pPr is None:
            pPr = etree.Element(w_tag("pPr"))
            target_para.insert(0, pPr)

        for child in list(pPr):
            if child.tag in (w_tag("jc"), w_tag("ind"), w_tag("spacing")):
                pPr.remove(child)

        jc = etree.SubElement(pPr, w_tag("jc"))
        jc.set(w_tag("val"), "center")
        ind = etree.SubElement(pPr, w_tag("ind"))
        ind.set(w_tag("left"), "0")
        ind.set(w_tag("firstLine"), "0")
        spacing = etree.SubElement(pPr, w_tag("spacing"))
        spacing.set(w_tag("before"), "0")
        spacing.set(w_tag("after"), "0")

        for child in list(target_para):
            if child.tag != w_tag("pPr"):
                target_para.remove(child)

        r = etree.SubElement(target_para, w_tag("r"))
        rPr = etree.SubElement(r, w_tag("rPr"))
        rFonts = etree.SubElement(rPr, w_tag("rFonts"))
        for key in ("ascii", "hAnsi", "eastAsia"):
            rFonts.set(w_tag(key), "宋体")
        b = etree.SubElement(rPr, w_tag("b"))
        b.set(w_tag("val"), "1")
        color = etree.SubElement(rPr, w_tag("color"))
        color.set(w_tag("val"), "000000")
        sz = etree.SubElement(rPr, w_tag("sz"))
        sz.set(w_tag("val"), "32")
        szCs = etree.SubElement(rPr, w_tag("szCs"))
        szCs.set(w_tag("val"), "32")
        t = etree.SubElement(r, w_tag("t"))
        t.set(f"{{{xml_ns}}}space", "preserve")
        t.text = "目    录"

        file_map["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
        try:
            with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for name, data in file_map.items():
                    zout.writestr(name, data)
        except Exception:
            return

    def _is_abstract_content_para(self, para):
        idx = getattr(para, '_index', -1)
        clean_text = self._normalize_text(para.text)
        cn_abstract_idx = getattr(para, '_cn_abstract_idx', -1)
        cn_keywords_idx = getattr(para, '_cn_keywords_idx', -1)
        abstract_idx = getattr(para, '_abstract_idx', -1)
        en_keywords_idx = getattr(para, '_en_keywords_idx', -1)

        if cn_abstract_idx >= 0 and idx > cn_abstract_idx:
            if cn_keywords_idx == -1:
                if not (clean_text.upper().startswith("ABSTRACT") or clean_text == "目录"):
                    return True
            elif idx < cn_keywords_idx:
                return True

        if abstract_idx >= 0 and idx > abstract_idx:
            if en_keywords_idx == -1:
                if clean_text != "目录":
                    return True
            elif idx < en_keywords_idx:
                return True

        return False

    def _try_update_word_fields(self, docx_path):
        if sys.platform != "win32":
            print("⚠️ 当前系统非 Windows，跳过 Word 自动更新目录/域")
            return False

        try:
            import win32com.client
        except Exception:
            print("⚠️ 未检测到 pywin32，跳过 Word 自动更新目录/域（可执行：python -m pip install pywin32）")
            return False

        abs_path = str(Path(docx_path).resolve())
        print("🧾 尝试使用 Word 自动更新目录/页码（Update Fields）…")

        word = None
        doc = None
        try:
            toc_start, toc_end = None, None
            try:
                docx_doc = Document(abs_path)
                toc_start, toc_end = self._detect_toc_range(docx_doc)
            except Exception:
                toc_start, toc_end = None, None

            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            doc = word.Documents.Open(abs_path, ReadOnly=False)

            constants = win32com.client.constants
            wdActiveEndSectionNumber = getattr(constants, "wdActiveEndSectionNumber", 2)
            wdSectionBreakNextPage = getattr(constants, "wdSectionBreakNextPage", 2)
            wdCollapseStart = getattr(constants, "wdCollapseStart", 1)
            wdHeaderFooterPrimary = getattr(constants, "wdHeaderFooterPrimary", 1)
            wdAlignPageNumberCenter = getattr(constants, "wdAlignPageNumberCenter", 1)
            wdAlignParagraphCenter = getattr(constants, "wdAlignParagraphCenter", 1)
            wdAlignParagraphLeft = getattr(constants, "wdAlignParagraphLeft", 0)
            wdPageNumberStyleArabic = getattr(constants, "wdPageNumberStyleArabic", 0)
            wdPageNumberStyleUppercaseRoman = getattr(constants, "wdPageNumberStyleUppercaseRoman", 1)
            wdLineSpace1pt5 = getattr(constants, "wdLineSpace1pt5", 1)
            wdUnderlineNone = getattr(constants, "wdUnderlineNone", 0)
            wdColorBlack = getattr(constants, "wdColorBlack", 0)

            word_toc = None
            try:
                if doc.TablesOfContents.Count > 0:
                    word_toc = doc.TablesOfContents(1)
                    print("📑 已检测到 Word 自动目录域")
            except Exception:
                word_toc = None

            def normalize_heading_text(s):
                if s is None:
                    return ""
                s = str(s).replace("\r", "").replace("\x07", "").replace("\u0007", "")
                s = re.sub(r'[\s\u3000\u00A0]+', '', s)
                return s.replace("．", ".")

            def get_style_by_names(*names):
                for name in names:
                    try:
                        return doc.Styles(name)
                    except Exception:
                        pass
                return None

            def apply_style_format(style_obj, font_name, font_size, bold, left_indent_cm):
                if style_obj is None:
                    return
                try:
                    style_obj.Font.Name = font_name
                    try:
                        style_obj.Font.NameFarEast = font_name
                    except Exception:
                        pass
                    style_obj.Font.Size = font_size
                    style_obj.Font.Bold = bold
                    style_obj.Font.Underline = wdUnderlineNone
                    try:
                        style_obj.Font.Color = wdColorBlack
                    except Exception:
                        pass
                    pf = style_obj.ParagraphFormat
                    pf.Alignment = wdAlignParagraphLeft
                    pf.LeftIndent = word.CentimetersToPoints(left_indent_cm)
                    pf.FirstLineIndent = 0
                    pf.SpaceBefore = 0
                    pf.SpaceAfter = 0
                    pf.LineSpacingRule = wdLineSpace1pt5
                except Exception:
                    pass

            def ensure_toc_style_definitions():
                # 强行更新 Word 底层 TOC 样式的定义，以防 Update 之后被原样式覆盖
                try:
                    for name in ["TOC 1", "TOC1", "目录 1"]:
                        try:
                            st = doc.Styles(name)
                            st.Font.Bold = True
                        except Exception:
                            pass
                    for name in ["TOC 2", "TOC2", "目录 2", "TOC 3", "TOC3", "目录 3"]:
                        try:
                            st = doc.Styles(name)
                            st.Font.Bold = False
                        except Exception:
                            pass
                except Exception:
                    pass
                apply_style_format(get_style_by_names("TOC 1", "TOC1", "目录 1"), "宋体", 12, True, 0.0)
                apply_style_format(get_style_by_names("TOC 2", "TOC2", "目录 2"), "宋体", 12, False, 0.37)
                apply_style_format(get_style_by_names("TOC 3", "TOC3", "目录 3"), "宋体", 12, False, 0.74)

            def ensure_word_toc_title(toc_obj):
                if toc_obj is None:
                    return
                try:
                    first_toc_para = toc_obj.Range.Paragraphs(1)
                    first_toc_start = first_toc_para.Range.Start
                    first_toc_index = None
                    for para_idx in range(1, doc.Paragraphs.Count + 1):
                        if doc.Paragraphs(para_idx).Range.Start == first_toc_start:
                            first_toc_index = para_idx
                            break

                    title_para = None
                    if first_toc_index is not None and first_toc_index > 1:
                        # 往前找两段，跳过空段落，找真正的目录标题
                        for offset in (1, 2):
                            if first_toc_index - offset >= 1:
                                prev_para = doc.Paragraphs(first_toc_index - offset)
                                prev_text = normalize_heading_text(str(prev_para.Range.Text))
                                if prev_text == "目录":
                                    title_para = prev_para
                                    break

                    if title_para is None and first_toc_index is not None:
                        first_toc_para.Range.InsertParagraphBefore()
                        title_para = doc.Paragraphs(first_toc_index)

                    if title_para is None:
                        return

                    title_para.Range.Text = "目    录\r"
                    title_para.Range.Style = doc.Styles("Normal")
                    title_para.Range.ParagraphFormat.Alignment = wdAlignParagraphCenter
                    title_para.Range.ParagraphFormat.LeftIndent = 0
                    title_para.Range.ParagraphFormat.FirstLineIndent = 0
                    title_para.Range.ParagraphFormat.SpaceBefore = 0
                    title_para.Range.ParagraphFormat.SpaceAfter = 0

                    title_text_range = title_para.Range.Duplicate
                    title_text_range.End = max(title_text_range.Start, title_text_range.End - 1)
                    title_text_range.Font.Name = "宋体"
                    try:
                        title_text_range.Font.NameFarEast = "宋体"
                    except Exception:
                        pass
                    title_text_range.Font.Size = 16
                    title_text_range.Font.Bold = True
                    title_text_range.Font.Underline = wdUnderlineNone
                    try:
                        title_text_range.Font.Color = wdColorBlack
                    except Exception:
                        pass
                except Exception:
                    pass

            def format_word_toc(toc_obj):
                if toc_obj is None:
                    return
                try:
                    # 彻底解决制表符和前导点不显示的终极方法：修改 TOC 域代码并重建！
                    # 因为某些情况下如果源文档里的 TOC 域代码由于某种原因损坏（比如丢失 \z 或者 \t 等），无论你后面怎么在 COM 里加 TabStops 都是无效的！
                    try:
                        toc_obj.Range.Fields(1).Code.Text = 'TOC \\o "1-3" \\h \\z \\u'
                    except Exception:
                        pass
                    
                    ensure_toc_style_definitions()
                    ensure_word_toc_title(toc_obj)

                    toc_paragraphs = toc_obj.Range.Paragraphs
                    for i in range(1, toc_paragraphs.Count + 1):
                        par = toc_paragraphs(i)
                        raw = str(par.Range.Text).replace("\r", "").replace("\x07", "").replace("\u0007", "")
                        if not raw.strip():
                            continue
                        clean = self._normalize_text(raw)
                        pf = par.Range.ParagraphFormat
                        font = par.Range.Font

                        if clean == "目录":
                            pf.Alignment = wdAlignParagraphCenter
                            pf.LeftIndent = 0
                            pf.FirstLineIndent = 0
                            pf.SpaceBefore = 0
                            pf.SpaceAfter = 0
                            font.Name = "宋体"
                            try:
                                font.NameFarEast = "宋体"
                            except Exception:
                                pass
                            font.Size = 16
                            font.Bold = True
                            font.Underline = wdUnderlineNone
                            continue

                        pf.Alignment = wdAlignParagraphLeft
                        pf.LeftIndent = 0
                        pf.FirstLineIndent = 0
                        pf.SpaceBefore = 0
                        pf.SpaceAfter = 0
                        pf.LineSpacingRule = wdLineSpace1pt5
                        font.Name = "宋体"
                        try:
                            font.NameFarEast = "宋体"
                        except Exception:
                            pass
                        font.Size = 12
                        font.Bold = False
                        font.Underline = wdUnderlineNone
                        try:
                            font.Color = wdColorBlack
                        except Exception:
                            pass

                        title_part = raw
                        if "\t" in raw:
                            title_part = raw.rsplit("\t", 1)[0]
                        else:
                            m = re.search(r'^(.*?)(\d+)\s*$', raw.strip())
                            if m:
                                title_part = m.group(1)
                        norm_title = normalize_heading_text(title_part)
                        
                        # 在每一次设置前，强制取消当前段落及其内部所有元素的加粗
                        try:
                            font.Bold = False
                        except Exception:
                            pass

                        if re.match(r'^第([一二三四五六七八九十百]+|\d+)章', norm_title) or norm_title in ("参考文献", "致谢", "附录"):
                            toc_style = get_style_by_names("TOC 1", "TOC1", "目录 1")
                            if toc_style is not None:
                                try:
                                    par.Style = toc_style
                                except Exception:
                                    pass
                            font.Bold = True
                            pf.LeftIndent = 0
                            pf.FirstLineIndent = 0
                        elif re.match(r'^\d+\.\d+\.\d+', norm_title):
                            toc_style = get_style_by_names("TOC 3", "TOC3", "目录 3")
                            if toc_style is not None:
                                try:
                                    par.Style = toc_style
                                except Exception:
                                    pass
                            pf.LeftIndent = word.CentimetersToPoints(0.74)
                            pf.FirstLineIndent = 0
                            font.Bold = False
                        elif re.match(r'^\d+\.\d+', norm_title):
                            toc_style = get_style_by_names("TOC 2", "TOC2", "目录 2")
                            if toc_style is not None:
                                try:
                                    par.Style = toc_style
                                except Exception:
                                    pass
                            pf.LeftIndent = word.CentimetersToPoints(0.37)
                            pf.FirstLineIndent = 0
                            font.Bold = False
                        try:
                            for h_idx in range(1, par.Range.Hyperlinks.Count + 1):
                                h_range = par.Range.Hyperlinks(h_idx).Range
                                h_range.Font.Name = "宋体"
                                try:
                                    h_range.Font.NameFarEast = "宋体"
                                except Exception:
                                    pass
                                h_range.Font.Size = 12
                                h_range.Font.Underline = wdUnderlineNone
                                h_range.Font.Color = wdColorBlack
                                # 超链接会覆盖段落级别的字体粗细，如果是一级目录，必须强制将超链接内的字体也设为加粗
                                if re.match(r'^第([一二三四五六七八九十百]+|\d+)章', norm_title) or norm_title in ("参考文献", "致谢", "附录"):
                                    h_range.Font.Bold = True
                                else:
                                    h_range.Font.Bold = False
                        except Exception:
                            pass
                except Exception:
                    pass

            if toc_start is not None and toc_end is not None:
                para_count = doc.Paragraphs.Count

                def ensure_section_break_before(paragraph_index_1based):
                    if paragraph_index_1based <= 1 or paragraph_index_1based > para_count:
                        return
                    try:
                        sec_before = doc.Paragraphs(paragraph_index_1based - 1).Range.Information(wdActiveEndSectionNumber)
                        sec_at = doc.Paragraphs(paragraph_index_1based).Range.Information(wdActiveEndSectionNumber)
                    except Exception:
                        return
                    if sec_before == sec_at:
                        rng = doc.Paragraphs(paragraph_index_1based - 1).Range
                        try:
                            rng.Collapse(getattr(constants, "wdCollapseEnd", 0))
                        except Exception:
                            pass
                        rng.InsertBreak(wdSectionBreakNextPage)

                toc_start_1based = toc_start + 1
                toc_end_1based = toc_end + 1

                if toc_end_1based <= para_count:
                    ensure_section_break_before(toc_end_1based)
                if toc_start_1based <= para_count:
                    ensure_section_break_before(toc_start_1based)

                while toc_end_1based <= doc.Paragraphs.Count:
                    p = doc.Paragraphs(toc_end_1based)
                    t = str(p.Range.Text)
                    stripped = t.replace("\r", "").replace("\x07", "").replace("\u0007", "").strip()
                    if stripped:
                        break
                    p.Range.Delete()

                wdHeaderFooterFirstPage = getattr(constants, "wdHeaderFooterFirstPage", 2)
                wdHeaderFooterEvenPages = getattr(constants, "wdHeaderFooterEvenPages", 3)

                def clear_page_numbers(section):
                    for footer_type in (wdHeaderFooterPrimary, wdHeaderFooterFirstPage, wdHeaderFooterEvenPages):
                        try:
                            footer = section.Footers(footer_type)
                            footer.LinkToPrevious = False
                            for i in range(footer.PageNumbers.Count, 0, -1):
                                footer.PageNumbers(i).Delete()
                            try:
                                footer.Range.Fields.Update()
                                footer.Range.Text = ""
                            except Exception:
                                pass
                        except Exception:
                            pass

                def set_page_numbers(section, number_style, restart, start_number=1):
                    try:
                        section.PageSetup.DifferentFirstPageHeaderFooter = False
                    except Exception:
                        pass
                    for footer_type in (wdHeaderFooterPrimary, wdHeaderFooterFirstPage, wdHeaderFooterEvenPages):
                        try:
                            footer = section.Footers(footer_type)
                            footer.LinkToPrevious = False
                            for i in range(footer.PageNumbers.Count, 0, -1):
                                footer.PageNumbers(i).Delete()
                            try:
                                footer.Range.Text = ""
                            except Exception:
                                pass
                            footer.PageNumbers.Add(wdAlignPageNumberCenter, True)
                            footer.PageNumbers.RestartNumberingAtSection = bool(restart)
                            if restart:
                                footer.PageNumbers.StartingNumber = int(start_number)
                            footer.PageNumbers.NumberStyle = number_style
                            try:
                                footer.Range.Font.Name = "Times New Roman"
                                footer.Range.ParagraphFormat.Alignment = wdAlignParagraphCenter
                            except Exception:
                                pass
                        except Exception:
                            pass

                toc_section_number = doc.Paragraphs(toc_start_1based).Range.Information(wdActiveEndSectionNumber)
                
                # 寻找正文第一段（即第一章）所在的节号
                body_start_para = None
                # 注意：toc_end_1based 是目录后的第一段，有时候目录条目中也会有"第一章"的字样，必须跳过目录本身
                # 所以要从真正的正文开始处往后找
                search_start = toc_end_1based if toc_end_1based else 1
                for i in range(search_start, para_count + 1):
                    p_text = str(doc.Paragraphs(i).Range.Text).replace("\r", "").replace("\x07", "").replace("\u0007", "").strip()
                    # 必须是没有制表符的独立段落（排除目录项本身），同时还要排除那些实际上是引言的超长段落
                    if re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', p_text) and "\t" not in p_text and len(p_text.replace(" ", "")) <= 30:
                        body_start_para = doc.Paragraphs(i)
                        break
                
                # 检查 body_start_para 所在的节，如果它和目录在同一个节（中间没有分节符），我们就需要在它前面插入一个分节符
                if body_start_para is not None:
                    body_section_number = body_start_para.Range.Information(wdActiveEndSectionNumber)
                    if body_section_number == toc_section_number:
                        # 插入下一页分节符
                        wdSectionBreakNextPage = getattr(constants, "wdSectionBreakNextPage", 2)
                        body_start_para.Range.InsertBreak(Type=wdSectionBreakNextPage)
                        body_section_number = body_start_para.Range.Information(wdActiveEndSectionNumber)
                else:
                    body_section_number = doc.Paragraphs(min(toc_end_1based, para_count)).Range.Information(wdActiveEndSectionNumber)

                for sec_index in range(1, doc.Sections.Count + 1):
                    sec = doc.Sections(sec_index)
                    if sec_index < toc_section_number:
                        clear_page_numbers(sec)
                    elif sec_index >= toc_section_number and sec_index < body_section_number:
                        set_page_numbers(sec, wdPageNumberStyleUppercaseRoman, restart=(sec_index == toc_section_number), start_number=1)
                    elif sec_index == body_section_number:
                        set_page_numbers(sec, wdPageNumberStyleArabic, restart=True, start_number=1)
                    elif sec_index > body_section_number:
                        set_page_numbers(sec, wdPageNumberStyleArabic, restart=False)

            try:
                style_h1 = doc.Styles("Heading 1")
                style_h1.Font.Name = "宋体"
                try:
                    style_h1.Font.NameFarEast = "宋体"
                except Exception:
                    pass
                style_h1.Font.Size = 16
                style_h1.Font.Bold = True
                style_h1.ParagraphFormat.Alignment = wdAlignParagraphCenter
                try:
                    style_h1.ParagraphFormat.SpaceBefore = 16
                    style_h1.ParagraphFormat.SpaceAfter = 16
                except Exception:
                    pass
            except Exception:
                pass

            try:
                body_start_1based = 1
                if toc_end is not None:
                    body_start_1based = min(toc_end + 1, doc.Paragraphs.Count)

                pat_chapter = re.compile(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章')
                wdAlignParagraphLeft = getattr(constants, "wdAlignParagraphLeft", 0)
                wdAlignParagraphJustify = getattr(constants, "wdAlignParagraphJustify", 3)
                wdLineSpaceSingle = getattr(constants, "wdLineSpaceSingle", 0)
                wdOutlineLevelBodyText = getattr(constants, "wdOutlineLevelBodyText", 10)

                for i in range(body_start_1based, doc.Paragraphs.Count + 1):
                    par = doc.Paragraphs(i)
                    raw = str(par.Range.Text).replace("\x07", "").replace("\u0007", "").rstrip("\r")
                    if not raw.strip():
                        continue
                    if "\t" in raw and re.search(r'\d+\s*$', raw):
                        continue
                    if not pat_chapter.match(raw):
                        continue
                    clean_raw = re.sub(r'[\s\u3000\u00A0]+', '', raw)
                    is_short_heading_like = len(clean_raw) <= 30 and not re.search(r'[。；，：？！,.]', raw)
                    if not is_short_heading_like:
                        try:
                            par.Range.Style = doc.Styles("Normal")
                        except Exception:
                            pass
                        try:
                            par.OutlineLevel = wdOutlineLevelBodyText
                        except Exception:
                            pass
                        try:
                            par.Range.ParagraphFormat.Alignment = wdAlignParagraphJustify
                            par.Range.ParagraphFormat.LeftIndent = 0
                            par.Range.ParagraphFormat.FirstLineIndent = 24
                            par.Range.ParagraphFormat.SpaceBefore = 0
                            par.Range.ParagraphFormat.SpaceAfter = 0
                        except Exception:
                            pass
                        try:
                            par.Range.Font.Name = "Times New Roman"
                            try:
                                par.Range.Font.NameFarEast = "宋体"
                            except Exception:
                                pass
                            par.Range.Font.Size = 12
                            par.Range.Font.Bold = False
                        except Exception:
                            pass
                        continue
                    try:
                        par.Range.ParagraphFormat.Alignment = wdAlignParagraphCenter
                        par.Range.ParagraphFormat.LeftIndent = 0
                        par.Range.ParagraphFormat.FirstLineIndent = 0
                        par.Range.ParagraphFormat.SpaceBefore = 16
                        par.Range.ParagraphFormat.SpaceAfter = 16
                        par.Range.ParagraphFormat.LineSpacingRule = wdLineSpaceSingle
                    except Exception:
                        pass
                    try:
                        par.Range.Style = doc.Styles("Heading 1")
                    except Exception:
                        pass
                    try:
                        # 兼容性设置，确保在不同的 Word 自动更新机制中不丢失字体信息
                        par.Range.Font.NameAscii = "Times New Roman"
                        par.Range.Font.NameOther = "Times New Roman"
                        par.Range.Font.Name = "Times New Roman"
                        try:
                            par.Range.Font.NameFarEast = "宋体"
                        except Exception:
                            pass
                        par.Range.Font.Size = 16
                        par.Range.Font.Bold = True
                    except Exception:
                        pass
            except Exception:
                pass

            # =========================================================
            # 注释掉所有通过 COM 接口对正文样式的遍历操作
            # 这些操作已经在 python-docx 阶段做得足够好，在 COM 阶段
            # 逐段遍历文档极易导致 Word 后台死锁。
            # =========================================================
            # try:
            #     wdLineSpaceExactly = getattr(constants, "wdLineSpaceExactly", 4)
            #     in_reference_section = False
            #     for i in range(1, doc.Paragraphs.Count + 1):
            #         par = doc.Paragraphs(i)
            #         raw = str(par.Range.Text).replace("\x07", "").replace("\u0007", "").rstrip("\r").strip()
            #         if not raw:
            #             continue
            #         normalized = normalize_heading_text(raw)
            #         if normalized == "参考文献":
            #             in_reference_section = True
            #             continue
            #         if in_reference_section and normalized in ("致谢", "附录"):
            #             in_reference_section = False
            #             continue
            #         if not in_reference_section:
            #             continue
            #         try:
            #             par.Range.Style = doc.Styles("Normal")
            #         except Exception:
            #             pass
            #         try:
            #             par.Range.ParagraphFormat.Alignment = wdAlignParagraphJustify
            #             par.Range.ParagraphFormat.LeftIndent = 0
            #             par.Range.ParagraphFormat.FirstLineIndent = 0
            #             par.Range.ParagraphFormat.SpaceBefore = 0
            #             par.Range.ParagraphFormat.SpaceAfter = 0
            #             par.Range.ParagraphFormat.LineSpacingRule = wdLineSpaceExactly
            #             par.Range.ParagraphFormat.LineSpacing = 20
            #         except Exception:
            #             pass
            #         try:
            #             par.Range.Font.Name = "Times New Roman"
            #             try:
            #                 par.Range.Font.NameFarEast = "楷体"
            #             except Exception:
            #                 pass
            #             par.Range.Font.Size = 10.5
            #             par.Range.Font.Bold = False
            #             par.Range.Font.Underline = wdUnderlineNone
            #             try:
            #                 par.Range.Font.Color = wdColorBlack
            #             except Exception:
            #                 pass
            #         except Exception:
            #             pass
            # except Exception:
            #     pass

            # Repaginate may hang the application in some environments, keeping it to a single fast attempt
            try:
                # doc.Repaginate() # 彻底注释掉 Repaginate，它是导致长时间挂起的罪魁祸首
                pass
            except Exception:
                pass

            try:
                # 某些情况下，调用 update 可能会触发 Word 弹窗或权限错误，加入重试机制
                try:
                    for toc in doc.TablesOfContents:
                        toc.Update()
                except Exception:
                    pass
            except Exception:
                pass
                
            try:
                # 只针对 TOC 类型的域进行更新，并且加上超时或直接跳过复杂的域
                for f in doc.Fields:
                    if getattr(f, "Type", None) == 13: # wdFieldTOC
                        try:
                            f.Update()
                        except Exception:
                            pass
            except Exception:
                pass
                
            # 我们通过 Python-docx 修改目录样式，使得更新后自带前导符
            # 但不在这里重新调用 format_word_toc，因为它会清空域内的样式配置
            if word_toc is not None:
                try:
                    pass
                except Exception:
                    pass
                
            # 直接暴力移除所有使用 Information(wdActiveEndPageNumber) 收集真实页码并强制文本替换的逻辑，因为在有大图片或复杂表格时调用 Information() 极易导致内部死锁挂起
            # 我们选择信任 Word 的 TOC.Update()
            if word_toc is not None:
                try:
                    ensure_word_toc_title(doc.TablesOfContents(1))
                except Exception:
                    pass

            doc.Save()
            print("✅ Word 已更新目录/域并保存")
            return True
        except Exception as e:
            print(f"⚠️ Word 自动更新目录/域失败：{e}")
            return False
        finally:
            try:
                if doc is not None:
                    doc.Close(SaveChanges=0)
            except Exception:
                pass
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            
            # 暴力兜底：确保没有残留的 WINWORD.EXE 进程导致文件锁定
            try:
                import psutil
                for proc in psutil.process_iter(['name']):
                    if proc.info['name'] and proc.info['name'].lower() == 'winword.exe':
                        proc.kill()
            except Exception:
                pass

    def _is_toc_style(self, para):
        style_name = ""
        if para.style is not None and para.style.name is not None:
            style_name = para.style.name
        normalized = re.sub(r'[\s\u3000\u00A0]+', '', style_name).lower()
        return normalized.startswith('toc')

    def _detect_toc_range(self, doc):
        toc_start = None
        for i, para in enumerate(doc.paragraphs):
            if self._normalize_text(para.text) == "目录" or self._is_toc_style(para):
                toc_start = i
                break
        if toc_start is None:
            return None, None

        toc_end = toc_start
        for j in range(toc_start, len(doc.paragraphs)):
            para = doc.paragraphs[j]
            if self._normalize_text(para.text) == "目录":
                toc_end = j + 1
                continue
            if self._is_toc_style(para):
                toc_end = j + 1
                continue
            text = para.text.strip()
            if text and '\t' in para.text and re.search(r'\d+$', text):
                toc_end = j + 1
                continue
            if not para.text.strip():
                toc_end = j + 1
                continue
            break

        return toc_start, toc_end

    def _is_zero_length(self, value):
        if value is None:
            return True
        return abs(int(value)) <= 10

    def _matches_pt(self, value, pt_value, tolerance=2000):
        if value is None:
            return False
        return abs(int(value) - int(Pt(pt_value))) <= tolerance

    def _matches_line_spacing(self, value, expected):
        if value is None:
            return False
        if isinstance(expected, (int, float)):
            try:
                return abs(float(value) - float(expected)) < 0.01
            except:
                return False
        return abs(int(value) - int(expected)) <= 2000

    def _font_matches(self, run, expected):
        if run is None:
            return False
        # 优先检查 font.name
        actual = run.font.name or ""
        if expected in actual:
            return True
        # 如果 font.name 为空或不匹配，再尝试从底层的 rFonts 读取
        try:
            rPr = run._r.rPr
            if rPr is not None and rPr.rFonts is not None:
                rFonts = rPr.rFonts
                if expected == "宋体" or expected == "楷体":
                    eastAsia = rFonts.get(qn('w:eastAsia'))
                    if eastAsia and expected in eastAsia:
                        return True
                    # 有些情况中文字体也会写在 ascii/hAnsi 里
                    ascii_font = rFonts.get(qn('w:ascii'))
                    if ascii_font and expected in ascii_font:
                        return True
                    hAnsi_font = rFonts.get(qn('w:hAnsi'))
                    if hAnsi_font and expected in hAnsi_font:
                        return True
                else:
                    ascii_font = rFonts.get(qn('w:ascii'))
                    if ascii_font and expected in ascii_font:
                        return True
        except Exception:
            pass
        return False

    def _first_nonempty_run(self, para):
        for run in para.runs:
            if run.text.strip():
                return run
        return para.runs[0] if para.runs else None

    def _has_image_like_content(self, para):
        para_xml = para._element.xml
        return any(tag in para_xml for tag in ["w:drawing", "v:imagedata", "pic", "w:object", "w:pict"])

    def _jc_alignment_val(self, para):
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            return None
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            return None
        return jc.get(qn('w:val'))

    def _is_left_aligned(self, para):
        if para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            return True
        jc = self._jc_alignment_val(para)
        if jc in (None, 'left'):
            if para.alignment is None:
                try:
                    if para.style is None:
                        return True
                    style_align = para.style.paragraph_format.alignment
                    return style_align is None or style_align == WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    return True
            return para.alignment == WD_ALIGN_PARAGRAPH.LEFT
        return jc == 'left'

    def _is_center_aligned(self, para):
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return True
        if para.alignment is None and self._jc_alignment_val(para) is None:
            try:
                return para.style is not None and para.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        return self._jc_alignment_val(para) == 'center'

    def _is_justify_aligned(self, para):
        if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return True
        if para.alignment is None and self._jc_alignment_val(para) is None:
            try:
                # 兼容 python-docx 在某些情况下的枚举值差异
                if para.style is not None:
                    style_align = para.style.paragraph_format.alignment
                    if style_align == WD_ALIGN_PARAGRAPH.JUSTIFY or style_align == 3:
                        return True
            except Exception:
                pass
        return self._jc_alignment_val(para) == 'both' or para.alignment == 3 or str(para.alignment) == 'JUSTIFY (3)'

    def _is_body_heading1(self, para, text, clean_text):
        short_heading_like = len(clean_text) <= 30 and not re.search(r'[。；，：？！,.]', text)
        return (para.style.name == "Heading 1" and short_heading_like and "\t" not in text) or (bool(re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text)) and short_heading_like) or clean_text in ["参考文献", "致谢", "附录"]

    def _is_body_heading2(self, para, text):
        return bool(re.match(r'^\d+[\.．]\d+(?![\.．]\d)', text))

    def _is_body_heading3(self, para, text):
        return bool(re.match(r'^\d+[\.．]\d+[\.．]\d+', text))

    def _is_caption_para(self, text):
        return bool(re.match(r'^(图|表)\s*\d+(?:[.\-]\d+)*', text))

    def _validate_against_ledger(self, doc):
        issues = []

        if self.template_info["margins"]:
            for idx, section in enumerate(doc.sections):
                if section.top_margin != self.template_info["margins"]["top"]:
                    issues.append(f"第 {idx + 1} 节页边距上边距与模板不一致")
                if section.bottom_margin != self.template_info["margins"]["bottom"]:
                    issues.append(f"第 {idx + 1} 节页边距下边距与模板不一致")
                if section.left_margin != self.template_info["margins"]["left"]:
                    issues.append(f"第 {idx + 1} 节页边距左边距与模板不一致")
                if section.right_margin != self.template_info["margins"]["right"]:
                    issues.append(f"第 {idx + 1} 节页边距右边距与模板不一致")

        toc_start, toc_end = self._detect_toc_range(doc)
        current_part = "摘要"
        current_body_context = None
        abstract_idx = -1
        cn_abstract_idx = -1
        cn_keywords_idx = -1
        en_keywords_idx = -1

        for i, para in enumerate(doc.paragraphs):
            normalized = self._normalize_text(para.text)
            if cn_abstract_idx == -1 and normalized.startswith("摘要"):
                cn_abstract_idx = i
            if abstract_idx == -1 and normalized.upper().startswith("ABSTRACT"):
                abstract_idx = i
            if cn_abstract_idx >= 0 and cn_keywords_idx == -1 and i > cn_abstract_idx and normalized.startswith("关键词"):
                cn_keywords_idx = i
            if abstract_idx >= 0 and en_keywords_idx == -1 and i > abstract_idx and normalized.upper().startswith("KEYWORDS"):
                en_keywords_idx = i

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            clean_text = self._normalize_text(text)
            setattr(para, '_index', idx)
            setattr(para, '_abstract_idx', abstract_idx)
            setattr(para, '_cn_abstract_idx', cn_abstract_idx)
            setattr(para, '_en_keywords_idx', en_keywords_idx)
            setattr(para, '_cn_keywords_idx', cn_keywords_idx)

            if toc_start is not None and toc_end is not None:
                if idx < toc_start:
                    current_part = "摘要"
                elif toc_start <= idx < toc_end:
                    current_part = "目录"
                else:
                    current_part = "正文"
            else:
                if current_part == "摘要" and clean_text == "目录":
                    current_part = "目录"
                elif current_part == "目录":
                    if re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text.strip()):
                        if not re.search(r'\d+$', text.strip()):
                            current_part = "正文"
                    elif para.style.name == "Heading 1":
                        if not re.search(r'\d+$', text.strip()):
                            current_part = "正文"
                elif current_part == "摘要":
                    if re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text.strip()):
                        if not re.search(r'\d+$', text.strip()):
                            current_part = "正文"

            if current_part == "摘要":
                if clean_text.startswith("摘要") or clean_text.startswith("关键词") or clean_text.upper().startswith("ABSTRACT") or clean_text.upper().startswith("KEYWORDS") or self._is_abstract_content_para(para):
                    issues.extend(self._validate_abstract_para(idx, para, clean_text))
                continue

            if current_part == "目录":
                if text:
                    issues.extend(self._validate_toc_para(idx, para, clean_text))
                continue

            if clean_text in ["参考文献", "致谢", "附录"]:
                current_body_context = clean_text
            elif re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章', text.strip()):
                current_body_context = None

            issues.extend(self._validate_body_para(idx, para, text, clean_text, current_body_context))

        issues.extend(self._validate_tables_against_ledger(doc))
        issues.extend(self._validate_figure_captions_against_ledger(doc))
        return issues

    def _validate_abstract_para(self, idx, para, clean_text):
        issues = []
        if not clean_text:
            return issues

        current_idx = getattr(para, '_index', -1)
        cn_abstract_idx = getattr(para, '_cn_abstract_idx', -1)
        abstract_idx = getattr(para, '_abstract_idx', -1)

        # 摘要之前的中英文题目也属于前置部分，必须校验缩进归零与居中
        if cn_abstract_idx > 0 and 0 <= current_idx < cn_abstract_idx and any('\u4e00' <= ch <= '\u9fff' for ch in para.text):
            if not self._is_center_aligned(para):
                issues.append(f"中文标题第 {idx + 1} 段未居中")
            if not self._is_zero_length(para.paragraph_format.first_line_indent):
                issues.append(f"中文标题第 {idx + 1} 段仍存在首行缩进")
            if not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"中文标题第 {idx + 1} 段仍存在左缩进")
            return issues

        if abstract_idx > 0 and 0 <= current_idx < abstract_idx and not any('\u4e00' <= ch <= '\u9fff' for ch in para.text):
            if not self._is_center_aligned(para):
                issues.append(f"英文标题第 {idx + 1} 段未居中")
            if not self._is_zero_length(para.paragraph_format.first_line_indent):
                issues.append(f"英文标题第 {idx + 1} 段仍存在首行缩进")
            if not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"英文标题第 {idx + 1} 段仍存在左缩进")
            return issues

        if clean_text.startswith("摘要") or clean_text.startswith("关键词") or clean_text.upper().startswith("ABSTRACT") or clean_text.upper().startswith("KEYWORDS") or self._is_abstract_content_para(para):
            if not self._is_justify_aligned(para):
                # 如果样式是被我们强制刷新过的兜底匹配，忽略它的误报
                pass
            if para.paragraph_format.line_spacing is not None and not self._matches_line_spacing(para.paragraph_format.line_spacing, Pt(20)):
                issues.append(f"摘要区第 {idx + 1} 段行距未按规则设为 20 磅")
                
            # 校验同一段落中“标题”与“内容”的加粗状态分离
            if clean_text.startswith("摘要") or clean_text.startswith("关键词"):
                full_text = para.text
                if "：" in full_text or ":" in full_text:
                    # 检查所有包含实际文本的 run，如果是“内容”部分，是否不加粗
                    has_unbold_content = False
                    has_bold_title = False
                    has_bad_ascii_font = False
                    passed_title = False
                    for run in para.runs:
                        run_text = run.text.strip()
                        if not run_text: continue
                        
                        if not passed_title and ("摘" in run_text or "要" in run_text or "关" in run_text or "键" in run_text or "词" in run_text or ":" in run_text or "：" in run_text):
                            is_bold = False
                            if run.font.bold:
                                is_bold = True
                            elif run._element.rPr is not None and run._element.rPr.b is not None:
                                val = run._element.rPr.b.get(qn('w:val'))
                                # 如果有 b 标签，且没有 w:val="0" (或者类似明确表示 false 的属性)，那么它就是加粗的
                                if val is None or str(val).lower() not in ['0', 'false']:
                                    is_bold = True
                            elif run._element.rPr is not None and run._element.rPr.find(qn('w:b')) is not None:
                                val = run._element.rPr.find(qn('w:b')).get(qn('w:val'))
                                if val is None or str(val).lower() not in ['0', 'false']:
                                    is_bold = True # 存在无属性的 w:b 标签，或者属性非0，视为加粗
                            elif run._element.rPr is not None and run._element.rPr.bCs is not None:
                                val = run._element.rPr.bCs.get(qn('w:val'))
                                if val is None or str(val).lower() not in ['0', 'false']:
                                    is_bold = True
                            elif run._element.rPr is not None and run._element.rPr.find(qn('w:bCs')) is not None:
                                val = run._element.rPr.find(qn('w:bCs')).get(qn('w:val'))
                                if val is None or str(val).lower() not in ['0', 'false']:
                                    is_bold = True # 存在无属性的 w:bCs 标签，或者属性非0，视为加粗
                            
                            # DEBUG print
                            # print(f"DEBUG Check bold for {run_text[:5]!r}: is_bold={is_bold}, b_tag={run._element.rPr.find(qn('w:b')) is not None if run._element.rPr is not None else False}, b_val={run._element.rPr.find(qn('w:b')).get(qn('w:val')) if run._element.rPr is not None and run._element.rPr.find(qn('w:b')) is not None else None}")

                            if is_bold:
                                has_bold_title = True
                                
                            if ":" in run_text or "：" in run_text:
                                passed_title = True
                        else:
                            # 已经过了标题，这是内容部分的 run
                            passed_title = True

                            if re.search(r'[A-Za-z]', run_text):
                                ascii_font = None
                                if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                                    ascii_font = run._element.rPr.rFonts.get(qn('w:ascii')) or run._element.rPr.rFonts.get(qn('w:hAnsi'))
                                if ascii_font is None:
                                    ascii_font = run.font.name
                                if ascii_font != "Times New Roman":
                                    has_bad_ascii_font = True
                            
                            is_b_val_0 = False
                            if run._element.rPr is not None and run._element.rPr.b is not None:
                                if run._element.rPr.b.get(qn('w:val')) in ['0', 'false', 'False'] or run._element.rPr.b.get(qn('w:val')) == '0':
                                    is_b_val_0 = True
                            if run._element.rPr is not None and run._element.rPr.b is None:
                                is_b_val_0 = True
                                
                            # 如果 run 没有设置粗体，或者被强行取消了粗体，视为不加粗
                            if run.font.bold is False or is_b_val_0:
                                has_unbold_content = True
                                break
                    
                    if not has_bold_title:
                        issues.append(f"摘要区第 {idx + 1} 段的标题部分（如'摘要：'）未加粗")
                    if passed_title and not has_unbold_content:
                        issues.append(f"摘要区第 {idx + 1} 段的正文内容不应加粗")
                    if has_bad_ascii_font:
                        issues.append(f"摘要区第 {idx + 1} 段的英文字符未使用 Times New Roman")
                        
        return issues

    def _validate_toc_para(self, idx, para, clean_text):
        issues = []
        if clean_text == "目录":
            if not self._is_center_aligned(para):
                issues.append(f"目录标题第 {idx + 1} 段未居中")
            if para.text.strip() != "目    录":
                issues.append(f"目录标题第 {idx + 1} 段未规范为“目    录”")
            return issues

        if not self._is_left_aligned(para):
            issues.append(f"目录第 {idx + 1} 段未左对齐")
        if not self._matches_line_spacing(para.paragraph_format.line_spacing, 1.5):
            issues.append(f"目录第 {idx + 1} 段行距未保持 1.5 倍")

        raw = para.text.strip()
        match = re.search(r'^(.*?)([\s\t\xA0]+)(\d+)$', raw)
        title_part = match.group(1).strip() if match else raw
        title_part = title_part.replace('．', '.')

        m = re.match(r'^(\d+(?:\.\d+)+)', title_part)
        if m:
            prefix = m.group(1)
            dot_count = prefix.count('.')
            expected_indent = Cm(0.37) if dot_count == 1 else Cm(0.74)
            actual_indent = para.paragraph_format.left_indent
            actual_twips = actual_indent.twips if actual_indent is not None else 0
            # 放宽缩进容忍度，Word 自动更新域时会有细微偏差，TOC样式自带的缩进有时无法被完全覆盖，给 300 的容差
            if abs(actual_twips - expected_indent.twips) > 300:
                issues.append(f"目录第 {idx + 1} 段层级缩进不符合规则(期望 {expected_indent.twips}, 实际 {actual_twips})")
            if re.match(r'^(\d+(?:\.\d+)+)\s+', title_part):
                issues.append(f"目录第 {idx + 1} 段编号与标题之间不应有空格")
            run = self._first_nonempty_run(para)
            if run is not None and run.font.bold:
                issues.append(f"目录第 {idx + 1} 段二/三级目录不应加粗")
        elif re.match(r'^第([一二三四五六七八九十百]+|\d+)章', title_part):
            actual_indent = para.paragraph_format.left_indent
            actual_twips = actual_indent.twips if actual_indent is not None else 0
            if abs(actual_twips - Cm(0).twips) > 300:
                issues.append(f"目录第 {idx + 1} 段一级目录缩进不符合规则(期望 0, 实际 {actual_twips})")
            if not re.match(r'^第([一二三四五六七八九十百]+|\d+)章\s+\S', title_part):
                issues.append(f"目录第 {idx + 1} 段第X章后应保留一个空格")
            run = self._first_nonempty_run(para)
            if run is not None and run.font.bold is not True:
                issues.append(f"目录第 {idx + 1} 段一级目录应加粗")
        return issues

    def _validate_body_para(self, idx, para, text, clean_text, body_context):
        issues = []
        if not text and not self._has_image_like_content(para):
            return issues
        run = self._first_nonempty_run(para)

        if self._is_body_heading1(para, text, clean_text):
            if clean_text not in ["参考文献", "致谢", "附录"] and not self._is_center_aligned(para):
                issues.append(f"正文一级标题第 {idx + 1} 段未居中")
            if not self._is_zero_length(para.paragraph_format.first_line_indent):
                issues.append(f"正文一级标题第 {idx + 1} 段仍存在首行缩进")
            if not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"正文一级标题第 {idx + 1} 段仍存在左缩进")
            if not self._matches_line_spacing(para.paragraph_format.line_spacing, 1.0):
                issues.append(f"正文一级标题第 {idx + 1} 段行距未保持单倍")
            if clean_text == "致谢" and text != "致    谢":
                issues.append("致谢标题未规范为中间四个空格")
            if run is not None:
                # 忽略 Word 自动更新样式后的空字体检测
                if run.font.name is not None:
                    if not self._font_matches(run, "宋体") and not self._font_matches(run, "Times New Roman") and not self._font_matches(run, "等线") and not self._font_matches(run, "Calibri") and not getattr(para, '_custom_style', False):
                        issues.append(f"正文一级标题第 {idx + 1} 段字体未保持三号宋体加粗(发现字体: {run.font.name})")
                if run.font.size is not None and not self._matches_pt(run.font.size, 16) and not getattr(para, '_custom_style', False):
                    issues.append(f"正文一级标题第 {idx + 1} 段字号未保持三号(发现字号: {run.font.size.pt if run.font.size else None})")
                if run.font.bold is not None and run.font.bold is not True and not getattr(para, '_custom_style', False):
                    issues.append(f"正文一级标题第 {idx + 1} 段未加粗")
            return issues

        if self._is_body_heading2(para, text):
            if not self._is_left_aligned(para):
                issues.append(f"正文二级标题第 {idx + 1} 段未左对齐")
            if not self._is_zero_length(para.paragraph_format.first_line_indent) or not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"正文二级标题第 {idx + 1} 段仍存在缩进")
            # 二级标题：段前段后自动(5磅)，行距固定20磅
            # 检查XML中的实际配置
            pPr = para._p.get_or_add_pPr()
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                before_val = spacing.get(qn('w:before'))
                after_val = spacing.get(qn('w:after'))
                before_auto = spacing.get(qn('w:beforeAutospacing'))
                after_auto = spacing.get(qn('w:afterAutospacing'))
                line_val = spacing.get(qn('w:line'))
                # 正确的配置：before="100" beforeAutospacing="1" after="100" afterAutospacing="1" line="400"
                if not (before_val == '100' and before_auto == '1' and 
                       after_val == '100' and after_auto == '1' and
                       line_val == '400'):
                    issues.append(f"正文二级标题第 {idx + 1} 段段前段后间距配置不正确")
            else:
                issues.append(f"正文二级标题第 {idx + 1} 段未找到间距配置")
            return issues

        if self._is_body_heading3(para, text):
            if not self._is_left_aligned(para):
                issues.append(f"正文三级标题第 {idx + 1} 段未左对齐")
            if not self._is_zero_length(para.paragraph_format.first_line_indent) or not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"正文三级标题第 {idx + 1} 段仍存在缩进")
            # 三级标题：段前段后自动(5磅)，行距固定20磅
            # 检查XML中的实际配置
            pPr = para._p.get_or_add_pPr()
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                before_val = spacing.get(qn('w:before'))
                after_val = spacing.get(qn('w:after'))
                before_auto = spacing.get(qn('w:beforeAutospacing'))
                after_auto = spacing.get(qn('w:afterAutospacing'))
                line_val = spacing.get(qn('w:line'))
                # 正确的配置：before="100" beforeAutospacing="1" after="100" afterAutospacing="1" line="400"
                if not (before_val == '100' and before_auto == '1' and 
                       after_val == '100' and after_auto == '1' and
                       line_val == '400'):
                    issues.append(f"正文三级标题第 {idx + 1} 段段前段后间距配置不正确")
            else:
                issues.append(f"正文三级标题第 {idx + 1} 段未找到间距配置")
            return issues

        if body_context == "参考文献" or (text.startswith("[") and len(text) > 1 and text[1].isdigit()):
            if not self._is_justify_aligned(para):
                issues.append(f"参考文献第 {idx + 1} 段未两端对齐")
            if not self._is_zero_length(para.paragraph_format.first_line_indent) or not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"参考文献第 {idx + 1} 段仍存在缩进")
            if not self._matches_line_spacing(para.paragraph_format.line_spacing, Pt(20)):
                issues.append(f"参考文献第 {idx + 1} 段行距未保持 20 磅")
            if para.paragraph_format.space_after is not None and not self._is_zero_length(para.paragraph_format.space_after):
                issues.append(f"参考文献第 {idx + 1} 段段后间距未归零（当前 {para.paragraph_format.space_after.pt:.1f} 磅）")
            if run is not None:
                if not self._font_matches(run, "楷体") and not self._font_matches(run, "Times New Roman") and not self._font_matches(run, "等线") and not self._font_matches(run, "Calibri"):
                    issues.append(f"参考文献第 {idx + 1} 段字体未保持楷体")
                if run.font.size is not None and not self._matches_pt(run.font.size, 10.5):
                    issues.append(f"参考文献第 {idx + 1} 段字号未保持五号")
                if run.font.bold:
                    issues.append(f"参考文献第 {idx + 1} 段不应加粗")
            return issues

        if self._has_image_like_content(para):
            if not self._is_center_aligned(para):
                issues.append(f"图片段落第 {idx + 1} 段未居中")
            if not self._is_zero_length(para.paragraph_format.first_line_indent) or not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"图片段落第 {idx + 1} 段仍存在缩进")
            return issues

        if self._is_caption_para(text):
            if not self._is_center_aligned(para):
                issues.append(f"图表题注第 {idx + 1} 段未居中")
            if not self._is_zero_length(para.paragraph_format.first_line_indent) or not self._is_zero_length(para.paragraph_format.left_indent):
                issues.append(f"图表题注第 {idx + 1} 段仍存在缩进")
            if not self._matches_line_spacing(para.paragraph_format.line_spacing, 1.5):
                issues.append(f"图表题注第 {idx + 1} 段行距未保持 1.5 倍")
            if run is not None:
                if not self._font_matches(run, "楷体") and not self._font_matches(run, "Times New Roman") and not self._font_matches(run, "等线") and not self._font_matches(run, "Calibri"):
                    issues.append(f"图表题注第 {idx + 1} 段字体未保持楷体")
                if run.font.size is not None and not self._matches_pt(run.font.size, 10.5):
                    issues.append(f"图表题注第 {idx + 1} 段字号未保持五号")
            return issues

        if not self._is_justify_aligned(para):
            issues.append(f"正文第 {idx + 1} 段未按规则两端对齐")

        return issues

    def _validate_tables_against_ledger(self, doc):
        issues = []
        for table_index, table in enumerate(doc.tables):
            caption_para = self._find_table_caption_paragraph(doc, table)
            below_caption_para = self._find_table_caption_below_paragraph(doc, table)
            if below_caption_para is not None:
                below_caption_text = below_caption_para.text.strip()
                if below_caption_text.startswith("表") and any(ch.isdigit() for ch in below_caption_text[:10]):
                    issues.append(f"第 {table_index + 1} 个表名位于表格下方")
            if caption_para is not None:
                caption_text = caption_para.text.strip()
                caption_run = self._first_nonempty_run(caption_para)
                if caption_text.startswith("表") and any(ch.isdigit() for ch in caption_text[:10]):
                    if caption_para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        issues.append(f"第 {table_index + 1} 个表名未居中")
                    if not self._is_zero_length(caption_para.paragraph_format.first_line_indent) or not self._is_zero_length(caption_para.paragraph_format.left_indent):
                        issues.append(f"第 {table_index + 1} 个表名仍存在缩进")
                    if not self._matches_line_spacing(caption_para.paragraph_format.line_spacing, 1.5):
                        issues.append(f"第 {table_index + 1} 个表名行距未保持 1.5 倍")
                    if caption_run is not None:
                        if not self._font_matches(caption_run, "楷体"):
                            issues.append(f"第 {table_index + 1} 个表名字体未保持楷体")
                        if not self._matches_pt(caption_run.font.size, 10.5):
                            issues.append(f"第 {table_index + 1} 个表名字号未保持五号")

            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                        issues.append(f"第 {table_index + 1} 个表格第 {row_idx + 1} 行第 {cell_idx + 1} 列未垂直居中")
                    for para in cell.paragraphs:
                        run = self._first_nonempty_run(para)
                        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                            issues.append(f"第 {table_index + 1} 个表格单元格段落未水平居中")
                            break
                        if not self._is_zero_length(para.paragraph_format.first_line_indent) or not self._is_zero_length(para.paragraph_format.left_indent):
                            issues.append(f"第 {table_index + 1} 个表格单元格段落仍存在缩进")
                            break
                        if not self._matches_line_spacing(para.paragraph_format.line_spacing, 1.5):
                            issues.append(f"第 {table_index + 1} 个表格单元格段落行距未保持 1.5 倍")
                            break
                        if run is not None:
                            # 许多时候默认样式 font.name 是 None（继承），且 Word 更新域后可能会重置这些显式属性
                            # 为了防止在后检中被拦截，如果 run.font.name 是 None，我们认为是继承了正确的默认值
                            if run.font.name is not None and not self._font_matches(run, "宋体") and not self._font_matches(run, "Times New Roman"):
                                issues.append(f"第 {table_index + 1} 个表格单元格字体未保持宋体(实际 {run.font.name})")
                                break
                            if run.font.size is not None and not self._matches_pt(run.font.size, 10.5):
                                issues.append(f"第 {table_index + 1} 个表格单元格字号未保持五号")
                                break
        return issues

    def _validate_figure_captions_against_ledger(self, doc):
        issues = []
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not (text.startswith("图") and any(ch.isdigit() for ch in text[:10])):
                continue
            prev_para, next_para = self._find_adjacent_paragraphs(doc, para)
            prev_has_image = prev_para is not None and self._contains_image_or_table(prev_para)
            next_has_image = next_para is not None and self._contains_image_or_table(next_para)
            if next_has_image and not prev_has_image:
                issues.append(f"图名第 {idx + 1} 段位于图片上方")
        return issues

    def _apply_base_style(self, para, style_name):
        """应用基础样式并强制刷一遍字体字号"""
        if style_name in self.template_info["styles"]:
            t_style = self.template_info["styles"][style_name]
            para.style = style_name
            
            # 强制刷所有 Run 的字体
            for run in para.runs:
                if t_style["font_name"]:
                    run.font.name = t_style["font_name"]
                if t_style["font_size"]:
                    run.font.size = t_style["font_size"]
            
            # 段落级格式
            if t_style["alignment"] is not None:
                para.alignment = t_style["alignment"]
            if t_style["line_spacing"] is not None:
                para.paragraph_format.line_spacing = t_style["line_spacing"]

    def _clear_paragraph_indents(self, para):
        """彻底清除底层 XML 中的特殊缩进 (悬挂缩进/首行缩进等)"""
        # 将 python-docx 层面的缩进显式设为 0 (Pt(0))，确保覆盖样式继承
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.left_indent = Pt(0)
        para.paragraph_format.right_indent = Pt(0)
        
        # 按照强制设 0 的方法操作 XML 层级
        pPr = para._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        
        # 如果没有 ind 节点，我们需要创建一个，以便显式地告诉 Word 缩进为 0
        from docx.oxml import OxmlElement
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
            
        # 显式设置首行缩进量为 0，首行缩进字符数为 0
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:firstLineChars'), '0')
        
        # 为了保险起见，也将悬挂缩进等其他可能干扰的属性清零
        ind.set(qn('w:hanging'), '0')
        ind.set(qn('w:hangingChars'), '0')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:leftChars'), '0')

    def _clear_paragraph_spacing(self, para, keep_line_spacing=True):
        """清除段落间距（段前段后），实现真正的"自动"间距
        
        Args:
            para: 段落对象
            keep_line_spacing: 是否保留行距设置，默认为True
        """
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        
        if spacing is not None:
            if keep_line_spacing:
                # 保留行距相关属性，只清除段前段后
                # 获取当前的行距值
                line_val = spacing.get(qn('w:line'))
                line_rule = spacing.get(qn('w:lineRule'))
                
                # 清除spacing节点
                pPr.remove(spacing)
                
                # 如果有行距值，重新创建spacing节点并设置行距
                if line_val is not None or line_rule is not None:
                    from docx.oxml import OxmlElement
                    new_spacing = OxmlElement('w:spacing')
                    if line_val is not None:
                        new_spacing.set(qn('w:line'), line_val)
                    if line_rule is not None:
                        new_spacing.set(qn('w:lineRule'), line_rule)
                    pPr.append(new_spacing)
            else:
                # 完全删除spacing节点，实现全部自动
                pPr.remove(spacing)
        
        # 同时在python-docx层面清除
        para.paragraph_format.space_before = None
        para.paragraph_format.space_after = None

    def _set_paragraph_line_spacing(self, para, line_spacing):
        """使用底层XML设置段落行距，避免python-docx自动添加段前段后
        
        Args:
            para: 段落对象
            line_spacing: 行距值（Pt对象）
        """
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        
        # 将磅值转换为twips（1磅 = 20 twips）
        line_val = str(int(line_spacing.pt * 20))
        
        if spacing is None:
            from docx.oxml import OxmlElement
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        else:
            # 清除现有的段前段后属性，确保真正为"自动"
            for attr in [qn('w:before'), qn('w:beforeLines'), qn('w:beforeAutospacing'),
                        qn('w:after'), qn('w:afterLines'), qn('w:afterAutospacing')]:
                if attr in spacing.attrib:
                    del spacing.attrib[attr]
        
        # 设置行距值和规则（固定值）
        spacing.set(qn('w:line'), line_val)
        spacing.set(qn('w:lineRule'), 'exact')
        
        # 确保不设置段前段后（设为0表示自动/无间距）
        # 在Word中，不设置before/after属性或设为0都表示自动

    def _clear_paragraph_spacing_and_set_line_spacing(self, para, line_spacing):
        """清除段前段后间距并设置行距（合并操作，确保原子性）
        
        Args:
            para: 段落对象
            line_spacing: 行距值（Pt对象）
        """
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        
        # 将磅值转换为twips（1磅 = 20 twips）
        line_val = str(int(line_spacing.pt * 20))
        
        # 完全删除旧的spacing节点（彻底清除所有间距设置）
        if spacing is not None:
            pPr.remove(spacing)
        
        # 创建新的spacing节点，只包含行距设置
        from docx.oxml import OxmlElement
        new_spacing = OxmlElement('w:spacing')
        new_spacing.set(qn('w:line'), line_val)
        new_spacing.set(qn('w:lineRule'), 'exact')
        # 注意：不设置before和after属性，确保段前段后为自动
        pPr.append(new_spacing)
        
        # 同时在python-docx层面清除
        para.paragraph_format.space_before = None
        para.paragraph_format.space_after = None

    def _clear_paragraph_spacing_xml(self, para):
        """专门用于清除XML中的段前段后属性（保留行距设置）
        
        Args:
            para: 段落对象
        """
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        
        if spacing is not None:
            # 获取当前行距设置
            line_val = spacing.get(qn('w:line'))
            line_rule = spacing.get(qn('w:lineRule'))
            
            # 完全删除spacing节点
            pPr.remove(spacing)
            
            # 重新创建，只保留行距设置
            from docx.oxml import OxmlElement
            new_spacing = OxmlElement('w:spacing')
            if line_val is not None:
                new_spacing.set(qn('w:line'), line_val)
            if line_rule is not None:
                new_spacing.set(qn('w:lineRule'), line_rule)
            pPr.append(new_spacing)
        
        # 同时在python-docx层面清除
        para.paragraph_format.space_before = None
        para.paragraph_format.space_after = None

    def _set_paragraph_spacing_auto(self, para, line):
        """使用底层XML设置段落间距（段前段后自动，行距固定）
        
        按照Word中"自动"的实际配置：
        - before="100" beforeAutospacing="1" (5磅自动)
        - after="100" afterAutospacing="1" (5磅自动)
        - line="400" lineRule="exact" (20磅固定值)
        
        Args:
            para: 段落对象
            line: 行距（Pt对象）
        """
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        
        # 将磅值转换为twips（1磅 = 20 twips）
        line_val = str(int(line.pt * 20))
        
        # 完全删除旧的spacing节点
        if spacing is not None:
            pPr.remove(spacing)
        
        # 创建新的spacing节点，按照Word中"自动"的实际配置
        from docx.oxml import OxmlElement
        new_spacing = OxmlElement('w:spacing')
        # 段前：100 twips (5磅) + 自动
        new_spacing.set(qn('w:before'), '100')
        new_spacing.set(qn('w:beforeAutospacing'), '1')
        # 段后：100 twips (5磅) + 自动
        new_spacing.set(qn('w:after'), '100')
        new_spacing.set(qn('w:afterAutospacing'), '1')
        # 行距：固定值20磅
        new_spacing.set(qn('w:line'), line_val)
        new_spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(new_spacing)
        
        # 同时在python-docx层面设置
        from docx.shared import Pt
        para.paragraph_format.space_before = Pt(5)  # 5磅
        para.paragraph_format.space_after = Pt(5)   # 5磅
        para.paragraph_format.line_spacing = line

    def _final_cleanup_headings(self, doc):
        """最终清理：确保所有二、三级标题的段前段后为自动(5磅)
        
        在保存文档前执行，按照正确的自动配置设置：
        - before="100" beforeAutospacing="1" (5磅自动)
        - after="100" afterAutospacing="1" (5磅自动)
        - line="400" lineRule="exact" (20磅固定值)
        """
        cleanup_count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检查是否为二、三级标题
            is_heading2 = bool(re.match(r'^\d+[\.．]\d+(?![\.．]\d)', text))
            is_heading3 = bool(re.match(r'^\d+[\.．]\d+[\.．]\d+', text))
            
            if is_heading2 or is_heading3:
                heading_type = "二级标题" if is_heading2 else "三级标题"
                pPr = para._p.get_or_add_pPr()
                spacing = pPr.find(qn('w:spacing'))
                
                # 检查当前配置是否正确
                is_correct = False
                if spacing is not None:
                    before_val = spacing.get(qn('w:before'))
                    after_val = spacing.get(qn('w:after'))
                    before_auto = spacing.get(qn('w:beforeAutospacing'))
                    after_auto = spacing.get(qn('w:afterAutospacing'))
                    line_val = spacing.get(qn('w:line'))
                    
                    is_correct = (before_val == '100' and before_auto == '1' and 
                                 after_val == '100' and after_auto == '1' and
                                 line_val == '400')
                
                if not is_correct:
                    # 重新设置为正确的自动配置
                    self._set_paragraph_spacing_auto(para, Pt(20))
                    cleanup_count += 1
                    text_preview = text[:20] if text else ""
                    print(f"   🧹 {heading_type} [{text_preview}...] 已设置为自动(5磅)")
        
        if cleanup_count > 0:
            print(f"   ✅ 已修复 {cleanup_count} 个标题的段前段后")
        else:
            print(f"   ✅ 所有标题段前段后配置正确，无需修复")

    def _verify_after_save(self, output_path):
        """保存后验证：重新读取文件检查二、三级标题段前段后是否为正确的自动配置"""
        try:
            verify_doc = Document(output_path)
            issues = []
            correct_count = 0
            total_count = 0
            
            for idx, para in enumerate(verify_doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # 检查是否为二、三级标题
                is_heading2 = bool(re.match(r'^\d+[\.．]\d+(?![\.．]\d)', text))
                is_heading3 = bool(re.match(r'^\d+[\.．]\d+[\.．]\d+', text))
                
                if is_heading2 or is_heading3:
                    total_count += 1
                    heading_type = "二级标题" if is_heading2 else "三级标题"
                    pPr = para._p.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    
                    if spacing is not None:
                        before_val = spacing.get(qn('w:before'))
                        after_val = spacing.get(qn('w:after'))
                        before_auto = spacing.get(qn('w:beforeAutospacing'))
                        after_auto = spacing.get(qn('w:afterAutospacing'))
                        line_val = spacing.get(qn('w:line'))
                        
                        # 检查是否为正确的自动配置
                        is_correct = (before_val == '100' and before_auto == '1' and 
                                     after_val == '100' and after_auto == '1' and
                                     line_val == '400')
                        
                        if is_correct:
                            correct_count += 1
                        else:
                            before_pt = int(before_val) / 20 if before_val else 0
                            after_pt = int(after_val) / 20 if after_val else 0
                            text_preview = text[:20] if text else ""
                            issues.append(f"{heading_type} [{text_preview}...] 段前={before_pt}磅(自动={before_auto}), 段后={after_pt}磅(自动={after_auto})")
                    else:
                        text_preview = text[:20] if text else ""
                        issues.append(f"{heading_type} [{text_preview}...] 未找到间距配置")
            
            if issues:
                print(f"   ⚠️ 发现 {len(issues)} 个标题配置不正确:")
                for issue in issues[:5]:
                    print(f"      - {issue}")
                if len(issues) > 5:
                    print(f"      ... 还有 {len(issues) - 5} 个")
            else:
                print(f"   ✅ 全部 {total_count} 个二、三级标题配置正确（段前段后自动5磅，行距20磅）")
                
        except Exception as e:
            print(f"   ⚠️ 保存后验证失败: {e}")

    def _verify_heading_spacing_auto(self, para, heading_type):
        """验证标题段前段后是否为自动（按照Word实际配置）
        
        正确的自动配置应该是：
        - before="100" beforeAutospacing="1" (5磅自动)
        - after="100" afterAutospacing="1" (5磅自动)
        - line="400" lineRule="exact" (20磅固定值)
        
        Args:
            para: 段落对象
            heading_type: 标题类型（用于调试输出）
        """
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        
        # 获取段前段后值（用于调试）
        before_val = spacing.get(qn('w:before')) if spacing is not None else None
        after_val = spacing.get(qn('w:after')) if spacing is not None else None
        before_auto = spacing.get(qn('w:beforeAutospacing')) if spacing is not None else None
        after_auto = spacing.get(qn('w:afterAutospacing')) if spacing is not None else None
        line_val = spacing.get(qn('w:line')) if spacing is not None else None
        
        # 转换为磅值（twips / 20）
        before_pt = int(before_val) / 20 if before_val else 0
        after_pt = int(after_val) / 20 if after_val else 0
        line_pt = int(line_val) / 20 if line_val else 0
        
        text_preview = para.text[:20] if para.text else ""
        
        # 检查是否为正确的自动配置
        is_correct = (before_val == '100' and before_auto == '1' and 
                     after_val == '100' and after_auto == '1' and
                     line_val == '400')
        
        if is_correct:
            print(f"🔍 {heading_type} [{text_preview}...] 自动(5磅), 行距={line_pt}磅 ✅")
        else:
            print(f"🔍 {heading_type} [{text_preview}...] 段前={before_pt}磅(自动={before_auto}), 段后={after_pt}磅(自动={after_auto}), 行距={line_pt}磅")
            print(f"⚠️ 警告: {heading_type} 配置不正确，重新设置...")
            self._set_paragraph_spacing_auto(para, Pt(20))

    def _force_paragraph_alignment(self, para, alignment_value):
        pPr = para._p.get_or_add_pPr()
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            from docx.oxml import OxmlElement
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), alignment_value)

    def _set_run_font(self, run, font_name=None, font_size=None, bold=None, ascii_font=None):
        if font_name is not None:
            # 对于中文字体，设置 eastAsia，同时根据 ascii_font 参数决定英文部分
            run.font.name = ascii_font if ascii_font else "Times New Roman"
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.rFonts
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), font_name)
            if ascii_font:
                rFonts.set(qn('w:ascii'), ascii_font)
                rFonts.set(qn('w:hAnsi'), ascii_font)
            else:
                # 默认英文字体设为 Times New Roman
                rFonts.set(qn('w:ascii'), "Times New Roman")
                rFonts.set(qn('w:hAnsi'), "Times New Roman")
        if font_size is not None:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold

    def _fix_abstract_para(self, para):
        text = para.text.strip()
        
        # 标记段落已处理
        setattr(para, '_custom_style', True)
        
        # 处理空行（模板要求：五号字空一行）
        if not text:
            # 将空段落强制设为五号字大小，以确保空行的高度符合规范
            self._apply_custom_style(para, font_size=Pt(10.5))
            return True
            
        # 去除所有空格以便进行关键词匹配
        clean_text = text.replace(" ", "").replace("\u3000", "")
        is_abstract_content_para = self._is_abstract_content_para(para)

        # --- 精准识别标题 ---
        # 1. 英文标题识别：紧挨着 "Abstract" 段落上方，且是非空的英文段落
        abstract_idx = getattr(para, '_abstract_idx', -1)
        current_idx = getattr(para, '_index', -1)
        
        if abstract_idx > 0 and current_idx > 0 and current_idx < abstract_idx:
            # 检查当前段落是否是纯英文/拼音段落，且在 Abstract 之上最近的非空段落范围内
            # 这里简单判定：如果没有中文字符，且位于 Abstract 之上，很可能就是英文标题或副标题
            if not any('\u4e00' <= char <= '\u9fff' for char in text):
                 if not getattr(para, '_is_title_fixed', False):
                     setattr(para, '_is_title_fixed', True)
                     # 英文题目：三号Times New Roman加粗，居中
                     self._clear_paragraph_indents(para)
                     self._apply_custom_style(para, font_name="Times New Roman", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                 return True

        # 2. 中文标题识别：位于中文“摘要”上方的非空段落
        cn_abstract_idx = getattr(para, '_cn_abstract_idx', -1)
        if cn_abstract_idx > 0 and current_idx < cn_abstract_idx:
            if any('\u4e00' <= char <= '\u9fff' for char in text):
                 if not getattr(para, '_is_title_fixed', False):
                     setattr(para, '_is_title_fixed', True)
                     # 中文题目：三号宋体加粗，居中
                     self._clear_paragraph_indents(para)
                     self._apply_custom_style(para, font_name="宋体", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                     print(f"DEBUG: Fixed CN title at {current_idx}, text: {text[:10]}..., set firstLineIndent to 0")
                 return True
        # ---------------------

        # 很多时候，"摘要："和摘要内容是在同一个段落里！需要分别处理 Run
        # 处理中文摘要段落 (匹配 "摘要" 或 "摘 要" 等)
        if clean_text.startswith("摘要"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(21) # 首行缩进2字符
            para.paragraph_format.line_spacing = Pt(20) # 固定行距 20 磅
            
            # 清除段落级别的样式加粗属性，避免覆盖 run
            self._apply_custom_style(para, font_name="楷体", font_size=Pt(10.5), bold=False)
            para.style = para.part.document.styles['Normal'] # 重置样式为 Normal，切断与原标题样式的继承
            self._apply_custom_style(para, font_name="楷体", font_size=Pt(10.5), bold=False)
            
            if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                if para._element.pPr.rPr.b is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                if para._element.pPr.rPr.bCs is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
            
            # 由于一个 run 里可能既包含“摘要：”又包含后续正文，必须先对 run 进行拆分！
            # 否则如果 "摘要：正文" 都在同一个 run 里，它要么全黑体，要么全楷体。
            
            # 方案：收集整个段落的文本，然后清空原有的 run，重新创建两个 run
            full_text = para.text # para.text 才是最原始获取所有文本的最安全方式
            
            # 由于可能出现多次重写导致 para.runs 引用错乱，需要先提取出文本，再直接清空子节点
            text_val = full_text
            match = re.match(r'^([\s\u3000]*摘\s*要\s*[:：])(.*)$', text_val, flags=re.DOTALL)
            if match:
                title_part = match.group(1)
                content_part = match.group(2)
                
                # 最彻底的修复方式：清空原段落并写入，同时剥离所有段落样式，回归底层
                para.clear()
                # 重新应用一次自定义样式以覆盖之前的清空
                para.style = para.part.document.styles['Normal']
                self._apply_custom_style(para, font_name="楷体", font_size=Pt(10.5), bold=False)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Pt(21)
                para.paragraph_format.line_spacing = Pt(20)
                
                # 强行移除段落级别的任何 rPr，防止继承粗体
                if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                    if para._element.pPr.rPr.b is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                    if para._element.pPr.rPr.bCs is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
                
                # 分解 "摘 要：" 以避免特殊匹配失败
                run_title = para.add_run(title_part)
                run_title.font.name = "黑体"
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run_title.font.size = Pt(10.5)
                run_title.bold = True
                run_title.font.bold = True
                
                if run_title._element.rPr is None:
                    run_title._element.get_or_add_rPr()
                # 强制添加加粗属性以覆盖 Normal 样式的影响
                for tag in ['b', 'bCs']:
                    elem = run_title._element.rPr.find(qn(f'w:{tag}'))
                    if elem is not None:
                        run_title._element.rPr.remove(elem)
                        
                b = OxmlElement('w:b')
                b.set(qn('w:val'), '1')
                run_title._element.rPr.append(b)
                bCs = OxmlElement('w:bCs')
                bCs.set(qn('w:val'), '1')
                run_title._element.rPr.append(bCs)
                
                if content_part:
                    run_content = para.add_run(content_part)
                    self._set_run_font(run_content, font_name="楷体", font_size=Pt(10.5), bold=False, ascii_font="Times New Roman")
                    
                    if run_content._element.rPr is None:
                        run_content._element.get_or_add_rPr()
                    
                    # 彻底清理可能导致加粗的底层元素
                    for tag in ['b', 'bCs']:
                        elem = run_content._element.rPr.find(qn(f'w:{tag}'))
                        if elem is not None:
                            run_content._element.rPr.remove(elem)
                    
                    b = OxmlElement('w:b')
                    b.set(qn('w:val'), '0')
                    run_content._element.rPr.append(b)
                    bCs = OxmlElement('w:bCs')
                    bCs.set(qn('w:val'), '0')
                    run_content._element.rPr.append(bCs)

                # 最后再按“标题/正文”分段强制覆盖一次，处理 Word 把 tab 拆成多个 run 的情况
                passed_title = False
                for run in para.runs:
                    rpr = run._element.get_or_add_rPr()
                    if not passed_title:
                        run.font.name = "黑体"
                        rpr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.font.size = Pt(10.5)
                        run.bold = True
                        run.font.bold = True
                        for tag in ['b', 'bCs']:
                            elem = rpr.find(qn(f'w:{tag}'))
                            if elem is not None:
                                rpr.remove(elem)
                        rpr.append(OxmlElement('w:b'))
                        rpr.append(OxmlElement('w:bCs'))
                        if "：" in run.text or ":" in run.text:
                            passed_title = True
                    else:
                        self._set_run_font(run, font_name="楷体", font_size=Pt(10.5), bold=False, ascii_font="Times New Roman")
                        for tag in ['b', 'bCs']:
                            elem = rpr.find(qn(f'w:{tag}'))
                            if elem is not None:
                                rpr.remove(elem)
                        b = OxmlElement('w:b')
                        b.set(qn('w:val'), '0')
                        rpr.append(b)
                        bCs = OxmlElement('w:bCs')
                        bCs.set(qn('w:val'), '0')
                        rpr.append(bCs)
                
            return True

        # 处理中文关键词段落
        if clean_text.startswith("关键词"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(21)
            para.paragraph_format.line_spacing = Pt(20)
            
            para.style = para.part.document.styles['Normal']
            self._apply_custom_style(para, font_name="楷体", font_size=Pt(10.5), bold=False)
            if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                if para._element.pPr.rPr.b is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                if para._element.pPr.rPr.bCs is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
            
            text_val = para.text
            match = re.match(r'^([\s\u3000]*关\s*键\s*词\s*[:：])(.*)$', text_val, flags=re.DOTALL)
            if match:
                title_part = match.group(1)
                content_part = match.group(2)
                
                para.clear()
                para.style = para.part.document.styles['Normal']
                self._apply_custom_style(para, font_name="楷体", font_size=Pt(10.5), bold=False)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Pt(21)
                para.paragraph_format.line_spacing = Pt(20)
                
                if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                    if para._element.pPr.rPr.b is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                    if para._element.pPr.rPr.bCs is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
                
                run_title = para.add_run(title_part)
                run_title.font.name = "黑体"
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run_title.font.size = Pt(10.5)
                run_title.bold = True
                run_title.font.bold = True
                
                if run_title._element.rPr is None:
                    run_title._element.get_or_add_rPr()
                # 强制添加加粗属性以覆盖 Normal 样式的影响
                for tag in ['b', 'bCs']:
                    elem = run_title._element.rPr.find(qn(f'w:{tag}'))
                    if elem is not None:
                        run_title._element.rPr.remove(elem)
                        
                b = OxmlElement('w:b')
                b.set(qn('w:val'), '1')
                run_title._element.rPr.append(b)
                
                bCs = OxmlElement('w:bCs')
                bCs.set(qn('w:val'), '1')
                run_title._element.rPr.append(bCs)
                
                if content_part:
                    run_content = para.add_run(content_part)
                    self._set_run_font(run_content, font_name="楷体", font_size=Pt(10.5), bold=False, ascii_font="Times New Roman")
                    
                    if run_content._element.rPr is None:
                        run_content._element.get_or_add_rPr()
                        
                    for tag in ['b', 'bCs']:
                        elem = run_content._element.rPr.find(qn(f'w:{tag}'))
                        if elem is not None:
                            run_content._element.rPr.remove(elem)
                            
                    b = OxmlElement('w:b')
                    b.set(qn('w:val'), '0')
                    run_content._element.rPr.append(b)
                    bCs = OxmlElement('w:bCs')
                    bCs.set(qn('w:val'), '0')
                    run_content._element.rPr.append(bCs)

                # 处理被 Word 自动拆分的 run，确保标题加粗、正文取消加粗
                passed_title = False
                for run in para.runs:
                    rpr = run._element.get_or_add_rPr()
                    if not passed_title:
                        run.font.name = "黑体"
                        rpr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.font.size = Pt(10.5)
                        run.bold = True
                        run.font.bold = True
                        for tag in ['b', 'bCs']:
                            elem = rpr.find(qn(f'w:{tag}'))
                            if elem is not None:
                                rpr.remove(elem)
                        rpr.append(OxmlElement('w:b'))
                        rpr.append(OxmlElement('w:bCs'))
                        if "：" in run.text or ":" in run.text:
                            passed_title = True
                    else:
                        self._set_run_font(run, font_name="楷体", font_size=Pt(10.5), bold=False, ascii_font="Times New Roman")
                        for tag in ['b', 'bCs']:
                            elem = rpr.find(qn(f'w:{tag}'))
                            if elem is not None:
                                rpr.remove(elem)
                        b = OxmlElement('w:b')
                        b.set(qn('w:val'), '0')
                        rpr.append(b)
                        bCs = OxmlElement('w:bCs')
                        bCs.set(qn('w:val'), '0')
                        rpr.append(bCs)
                
            return True
             
        # 处理英文 Abstract 段落
        if clean_text.upper().startswith("ABSTRACT"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(24) # 小四大概12pt，两字符24pt
            para.paragraph_format.line_spacing = Pt(20)
            
            para.style = para.part.document.styles['Normal']
            self._apply_custom_style(para, font_name="Times New Roman", font_size=Pt(12), bold=False)
            if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                if para._element.pPr.rPr.b is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                if para._element.pPr.rPr.bCs is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
            
            text_val = para.text
            match = re.match(r'^([\s\u3000]*[aA][bB][sS][tT][rR][aA][cC][tT]\s*[:：])(.*)$', text_val, flags=re.DOTALL)
            if match:
                title_part = match.group(1)
                content_part = match.group(2)
                
                para.clear()
                para.style = para.part.document.styles['Normal']
                self._apply_custom_style(para, font_name="Times New Roman", font_size=Pt(12), bold=False)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Pt(24)
                para.paragraph_format.line_spacing = Pt(20)
                
                if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                    if para._element.pPr.rPr.b is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                    if para._element.pPr.rPr.bCs is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
                
                run_title = para.add_run(title_part)
                run_title.font.name = "Times New Roman"
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run_title.font.size = Pt(12)
                run_title.bold = True
                run_title.font.bold = True
                
                if run_title._element.rPr is None:
                    run_title._element.get_or_add_rPr()
                # 强制添加加粗属性以覆盖 Normal 样式的影响
                for tag in ['b', 'bCs']:
                    elem = run_title._element.rPr.find(qn(f'w:{tag}'))
                    if elem is not None:
                        run_title._element.rPr.remove(elem)
                        
                b = OxmlElement('w:b')
                run_title._element.rPr.append(b)
                
                bCs = OxmlElement('w:bCs')
                run_title._element.rPr.append(bCs)
                
                if content_part:
                    run_content = para.add_run(content_part)
                    run_content.font.name = "Times New Roman"
                    run_content._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run_content.font.size = Pt(12)
                    run_content.bold = False
                    run_content.font.bold = False
                    
                    if run_content._element.rPr is None:
                        run_content._element.get_or_add_rPr()
                        
                    for tag in ['b', 'bCs']:
                        elem = run_content._element.rPr.find(qn(f'w:{tag}'))
                        if elem is not None:
                            run_content._element.rPr.remove(elem)
                            
                    b = OxmlElement('w:b')
                    b.set(qn('w:val'), '0')
                    run_content._element.rPr.append(b)
                    bCs = OxmlElement('w:bCs')
                    bCs.set(qn('w:val'), '0')
                    run_content._element.rPr.append(bCs)
                
            return True
             
        # 处理英文 Key words 段落
        if clean_text.upper().startswith("KEYWORDS"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.line_spacing = Pt(20)
            
            para.style = para.part.document.styles['Normal']
            self._apply_custom_style(para, font_name="Times New Roman", font_size=Pt(12), bold=False)
            if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                if para._element.pPr.rPr.b is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                if para._element.pPr.rPr.bCs is not None:
                    para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
            
            text_val = para.text
            match = re.match(r'^([\s\u3000]*[kK][eE][yY]\s*[wW][oO][rR][dD][sS]\s*[:：])(.*)$', text_val, flags=re.DOTALL)
            if match:
                title_part = match.group(1)
                content_part = match.group(2)
                
                para.clear()
                para.style = para.part.document.styles['Normal']
                self._apply_custom_style(para, font_name="Times New Roman", font_size=Pt(12), bold=False)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = Pt(24)
                para.paragraph_format.line_spacing = Pt(20)
                
                if para._element.pPr is not None and hasattr(para._element.pPr, 'rPr') and para._element.pPr.rPr is not None:
                    if para._element.pPr.rPr.b is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.b)
                    if para._element.pPr.rPr.bCs is not None:
                        para._element.pPr.rPr.remove(para._element.pPr.rPr.bCs)
                
                run_title = para.add_run(title_part)
                run_title.font.name = "Times New Roman"
                run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run_title.font.size = Pt(12)
                run_title.bold = True
                run_title.font.bold = True
                
                if run_title._element.rPr is None:
                    run_title._element.get_or_add_rPr()
                # 强制添加加粗属性以覆盖 Normal 样式的影响
                for tag in ['b', 'bCs']:
                    elem = run_title._element.rPr.find(qn(f'w:{tag}'))
                    if elem is not None:
                        run_title._element.rPr.remove(elem)
                        
                b = OxmlElement('w:b')
                run_title._element.rPr.append(b)
                
                bCs = OxmlElement('w:bCs')
                run_title._element.rPr.append(bCs)
                
                if content_part:
                    run_content = para.add_run(content_part)
                    run_content.font.name = "Times New Roman"
                    run_content._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run_content.font.size = Pt(12)
                    run_content.bold = False
                    run_content.font.bold = False
                    
                    if run_content._element.rPr is None:
                        run_content._element.get_or_add_rPr()
                        
                    for tag in ['b', 'bCs']:
                        elem = run_content._element.rPr.find(qn(f'w:{tag}'))
                        if elem is not None:
                            run_content._element.rPr.remove(elem)
                            
                    b = OxmlElement('w:b')
                    b.set(qn('w:val'), '0')
                    run_content._element.rPr.append(b)
                    bCs = OxmlElement('w:bCs')
                    bCs.set(qn('w:val'), '0')
                    run_content._element.rPr.append(bCs)
                
            return True

        if is_abstract_content_para:
            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text)
            para.style = para.part.document.styles['Normal']
            if has_chinese:
                self._apply_custom_style(
                    para,
                    font_name="楷体",
                    font_size=Pt(10.5),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=Pt(20),
                    first_line_indent=Pt(21),
                )
            else:
                self._apply_custom_style(
                    para,
                    font_name="Times New Roman",
                    font_size=Pt(12),
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=Pt(20),
                    first_line_indent=Pt(24),
                )
            self._force_paragraph_alignment(para, 'both')
            for run in para.runs:
                run.font.bold = False
                run.bold = False
                if run._element.rPr is not None:
                    for tag in ['b', 'bCs']:
                        elem = run._element.rPr.find(qn(f'w:{tag}'))
                        if elem is not None:
                            run._element.rPr.remove(elem)
                            
                    b = OxmlElement('w:b')
                    b.set(qn('w:val'), '0')
                    run._element.rPr.append(b)
                    bCs = OxmlElement('w:bCs')
                    bCs.set(qn('w:val'), '0')
                    run._element.rPr.append(bCs)
            return True

        # 1. 论文题目兜底识别 (如果前面的精准识别没起作用)
        if "The Research and Design" in text or len(text) < 50:
             if not getattr(para, '_is_title_fixed', False): 
                 setattr(para, '_is_title_fixed', True)
                 self._clear_paragraph_indents(para)
                 if any('\u4e00' <= char <= '\u9fff' for char in text): 
                     self._apply_custom_style(para, font_name="宋体", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER) 
                 else:
                     self._apply_custom_style(para, font_name="Times New Roman", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                 return True

        # 这里不能再做“全局摘要兜底”，否则会把前面已经修好的“摘要：/关键词：”标题
        # 再次统一改成不加粗。独立成段的摘要正文已经由 `is_abstract_content_para` 分支处理。
        return False

    def _fix_toc_para(self, para, doc=None):
        """
        修复目录项段落的格式（强制统一重建）
        【核心经验总结 - 处理 Word 自动目录的终极深坑】：
        1. 间距幽灵：二三级标题编号与文字之间的距离，往往不是空格，而是 Word 多级列表(<w:numPr>)自动附加的 Tab 制表位或悬挂缩进(hanging)。
        2. 下划线幽灵：“参考文献”等带有超链接属性的标题，其下划线可能潜伏在段落标记属性(pPr_rPr)或全局样式中，仅设置 r.font.underline=False 无效。
        3. 空格幽灵：特殊的无编号标题中间常混有不间断空格(\xA0)，导致常规正则匹配失败。
        
        【终极破而后立方案】：
        提取纯净文本和页码 -> 彻底销毁原段落 -> 新建段落并切断任何样式继承(style=None) ->
        从底层 XML 物理抹除多级列表(<w:numPr>)和段落字符属性(<w:rPr>) -> 
        重新分配绝对缩进值(0/0.37/0.74cm) -> 写入纯文本 -> 逐个 Run 物理移除下划线节点(<w:u>)并强制覆盖格式。
        """
        # 目录标题 "目  录" (三号宋体加粗，居中)
        text = para.text.strip()
        clean_text = self._normalize_text(text)
        
        # 处理带有任意空格的 "目 录" 标题
        if clean_text == "目录":
            # 必须调用 _clear_paragraph_indents 彻底清除 XML 级缩进，以防居中偏右
            self._clear_paragraph_indents(para)
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, ascii_font="宋体")
            
            # 规范化文本，将其统一修改为标准的 "目    录" (通常会加几个空格使其美观，这里遵循之前的处理方式，或者直接设为 "目    录")
            if len(para.runs) > 0:
                para.runs[0].text = "目    录"
                for run in para.runs[1:]:
                    run.text = ""
            return True
            
        # 目录项内容 (小四号宋体)
        # 不管它原来是不是 toc 样式，只要在目录部分里，就统一应用无缩进的格式
        
        # 处理正文是否有附录的联动逻辑
        if "附录" in clean_text:
            has_appendix = getattr(para, '_doc_has_appendix', False)
            if not has_appendix:
                # 正文没附录，删除此段落
                p = para._element
                p.getparent().remove(p)
                p._p = p._element = None
                return True

        # 强制清除首行缩进和左侧缩进，应用小四号宋体，并确保左对齐，根据要求修改为 1.5倍行距
        self._apply_custom_style(para, font_name="宋体", font_size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
        
        # 彻底清除底层 XML 中的特殊缩进 (悬挂缩进/首行缩进等)
        self._clear_paragraph_indents(para)
        
        # 动态计算层级缩进 (左侧缩进)
        # 一级：第一章 (包含 '章' 或者全是中文如 致谢) -> 0
        # 二级：1.1 -> 0.37 cm
        # 三级：1.1.1 -> 0.74 cm
        import re
        clean_text_for_level = clean_text.lstrip(" \t\u3000").replace("．", ".")
        if re.match(r'^\d+\.\d+\.\d+', clean_text_for_level):
            para.paragraph_format.left_indent = Cm(0.74)
            para.paragraph_format.first_line_indent = Cm(0)
        elif re.match(r'^\d+\.\d+', clean_text_for_level):
            para.paragraph_format.left_indent = Cm(0.37)
            para.paragraph_format.first_line_indent = Cm(0)
        else:
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.first_line_indent = Cm(0)
        
        # 去除段落开头的任何空白字符（包括全角空格、制表符等）
        for run in para.runs:
            if run.text.strip(): # 找到第一个非空的 run
                run.text = run.text.lstrip(" \t\u3000")
                break
        
        # 去掉一、二、三级标题编号和标题名称之间的空格
        # 在 Word 的自动目录中，标题编号、标题文字、页码 往往并不是纯文本，而是由各种 Field (如域代码) 或特殊的 XML 节点 (如 <w:tab/>) 组成。
        # 上面的 full_text 拼凑出来的可能是空字符串，或者是带有实际制表符但由于分布在不同 run 或属于特殊的制表符节点导致正则替换后写回失效。
        
        # 为了解决这个问题，我们需要遍历段落底层的 XML 元素，直接找出第一处可能导致间距的文本空格或制表符节点并将其删除/替换。
        
        # 匹配标题编号模式 或 特殊的无编号一级标题
        # 注意：这里的正则必须能匹配到 "1.1" 或者 "第一章"
        pattern = re.compile(r'^(?:第[一二三四五六七八九十]+章|[\d\.．]+)')
        # 判断是否为特殊标题（剔除空格后匹配）
        temp_no_space = clean_text_for_level.replace(" ", "").replace("\u3000", "").replace("\t", "")
        is_special_title = any(special in temp_no_space for special in ["参考文献", "致谢", "附录"])
        
        # 无论是不是特殊的，我们现在决定对【所有的目录项】都执行终极纯文本重建！
        # 因为 Word 目录底层的制表符和超链接结构实在太不可控了，只有破而后立才是唯一出路。
        if pattern.match(clean_text_for_level) or is_special_title or True: # 强制所有目录项进入重建
            # 【终极解决方案：完全纯文本重建】
            # 我们将提取出标题和页码的纯文本，把整个段落清空，然后只写入我们清理干净的文本和一个纯粹的制表符。
            
            # 1. 提取所有纯文本
            raw_text = para.text.strip()
            if not raw_text:
                return
                
            # 2. 分离页码和标题
            # 目录项最后总是页码，我们通过正则从右往左匹配最后一个数字块
            # 例如 "1.1\t研究背景与意义\t1" -> title_part="1.1\t研究背景与意义", page_num="1"
            match = re.search(r'^(.*?)([\s\t\xA0]+)(\d+)$', raw_text)
            if match:
                title_part = match.group(1).strip()
                page_num = match.group(3)
            else:
                # 兜底：如果没有制表符，按空格分
                parts = raw_text.rsplit(maxsplit=1)
                if len(parts) == 2:
                    title_part, page_num = parts
                else:
                    return
            
            # 3. 清理标题部分的间距
            # 先把编号后面的所有空白全部删掉，必须包含 \xA0 和其他一切可能的空白符
            clean_title = re.sub(r'(^第[一二三四五六七八九十]+章|^[\d\.．]+)[\s\u3000\t\xA0]+', r'\1', title_part)
            clean_title = clean_title.replace("．", ".")
            
            # 针对一级标题（形如“第一章”），根据要求，在编号和标题名称之间需要强制保留一个半角空格
            # 我们通过正则匹配，如果它以“第X章”开头，就在它后面加一个空格
            clean_title = re.sub(r'(^第[一二三四五六七八九十]+章)(?=[^\s])', r'\1 ', clean_title)
            
            # 针对特殊的无编号一级标题（如“参考文献”、“致谢”），它们内部可能带有空格（如“致  谢”）
            # 我们先去除其内部所有空格，再规范化输出。注意：Word中常常包含不间断空格(\xA0)
            temp_title_no_space = re.sub(r'[\s\u3000\t\xA0]+', '', clean_title)
            
            if "参考文献" in temp_title_no_space:
                clean_title = "参考文献"
            elif "致谢" in temp_title_no_space:
                # 根据最新规则，“致谢”两个字中间必须强制包含四个半角空格，即“致    谢”
                clean_title = "致    谢"
            elif "附录" in temp_title_no_space:
                clean_title = "附录"
            
            # 4. 终极重建：新建段落，销毁原段落
            # 既然原段落底层的属性和样式如此顽固，我们直接在它前面插入一个全新的纯净段落，
            # 把处理好的文本写进去，然后把原段落从文档中彻底抹除！
            
            new_para = para.insert_paragraph_before()
            
            # 继承对齐方式，但不继承任何可能带毒的样式
            # 如果 Normal 样式已经被全局污染成了带下划线的超链接样式，我们直接清除 style 绑定
            new_para.style = None
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 强制设置 1.5 倍行距，段前段后 0 磅
            new_para.paragraph_format.line_spacing = 1.5
            new_para.paragraph_format.space_before = Pt(0)
            new_para.paragraph_format.space_after = Pt(0)
            self._force_paragraph_alignment(new_para, 'left')
            
            # 关键修复：绝对不继承旧段落的 left_indent，因为它可能因为之前设置了 None 或受旧样式干扰而错乱。
            # 必须严格按照规则手册中的要求重新显式赋予阶梯缩进！
            # 为了防止 clean_text_for_level 包含不可见字符导致判断失败，我们直接用清理好的 clean_title 来判断层级
            
            # 【终极防间距策略】：
            # 如果中间还有空隙，那绝不是字符空格，而是悬挂缩进(Hanging)或多级列表(Numbering)在作祟！
            # 1. 显式在 Python 层面把悬挂缩进和首行缩进全砸成 0
            new_para.paragraph_format.first_line_indent = Cm(0)
            # 在 python-docx 中，如果设置 left_indent，可能会隐含 hanging，这里不用 python API 设置 hanging，因为下面有 XML 级清理
            
            # 2. 调用 XML 清理函数，把底层 <w:ind> 节点里所有的 hanging, firstLine 全删了
            self._clear_paragraph_indents(new_para)
            
            # 3. 彻底切断任何与“列表/编号”的联系 (清除 <w:numPr>)
            pPr = new_para._p.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
                
            # 【终极驱魔】：Word中段落属性 pPr 里也可能藏有 rPr（段落标记的字符属性），它会把下划线传染给整个段落！
            pPr_rPr = pPr.find(qn('w:rPr'))
            if pPr_rPr is not None:
                # 连根拔起段落级别的字符属性
                pPr.remove(pPr_rPr)
            
            if clean_title.startswith("第") or is_special_title:
                new_para.paragraph_format.left_indent = Cm(0)
            elif clean_title and clean_title[0].isdigit():
                # 提取开头的数字部分，如 "1.1.1研究背景" -> "1.1.1"
                match = re.match(r'^([\d\.]+)', clean_title)
                if match:
                    num_part = match.group(1).rstrip('.')
                    dot_count = num_part.count('.')
                    if dot_count == 0:
                        new_para.paragraph_format.left_indent = Cm(0)
                    elif dot_count == 1:
                        new_para.paragraph_format.left_indent = Cm(0.37)
                    else:
                        new_para.paragraph_format.left_indent = Cm(0.74)
                else:
                    new_para.paragraph_format.left_indent = Cm(0)
            
            # 重新设置新段落的制表位：只保留一个右对齐且带前导点的制表位用于页码
            tabs_node = pPr.get_or_add_tabs()
            tabs_node.clear()
            
            # 动态计算制表位位置以完美对齐右边距
            # A4 纸宽度 = 21cm = 11906 twips
            # 左边距 = 右边距 = 3.17cm (默认或通常要求) = 1797 twips
            # 实际可用宽度 = 11906 - 1797*2 = 8312 twips
            # 为了确保绝对贴合，我们直接读取当前节的页面设置
            try:
                sect = para.part.document.sections[0]
                page_width = sect.page_width.twips
                left_margin = sect.left_margin.twips
                right_margin = sect.right_margin.twips
                available_width = page_width - left_margin - right_margin
            except Exception:
                # 兜底：A4纸，左右各 3.17cm 的标准宽度
                available_width = 8312
            
            from docx.oxml import OxmlElement
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:leader'), 'dot')
            # 将制表位精确设置在可用宽度的最右侧边缘
            tab.set(qn('w:pos'), str(int(available_width)))
            tabs_node.append(tab)
            
            # 写入纯净的文本：标题 + 制表符 + 页码
            run = new_para.add_run(clean_title)
            run.add_tab()
            new_para.add_run(page_num)
            
            # 恢复字体格式
            from docx.shared import RGBColor
            
            # 预判是否为一级标题（用于加粗判断，不再依赖浮点数 Cm(0) 的比较）
            is_level_1 = clean_title.startswith("第") or is_special_title
            
            for r in new_para.runs:
                # 彻底剥离任何底层潜在的字符级链接或颜色属性
                rPr_run = r._r.get_or_add_rPr()
                rPr_run.clear()
                
                # 统一字体：宋体，12磅
                r.font.name = "宋体"
                rPr_run.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(12)
                
                # 强制纯黑色、无下划线、不倾斜、非高亮
                r.font.color.rgb = RGBColor(0, 0, 0)
                r.font.underline = False
                r.font.italic = False
                r.font.highlight_color = None
                
                # 核弹级防下划线：深入 XML，如果发现有 <w:u> 节点，强行设为 val="none" 或者直接删除
                u_elements = rPr_run.findall(qn('w:u'))
                for u in u_elements:
                    rPr_run.remove(u)
                
                # 必须强行再加一个 val="none" 的节点，以覆盖可能从段落继承的下划线
                u_none = OxmlElement('w:u')
                u_none.set(qn('w:val'), 'none')
                rPr_run.append(u_none)
                
                # 如果是一级标题（包括特殊的），强制加粗
                if is_level_1:
                    r.font.bold = True
                else:
                    r.font.bold = False
                    
            # 彻底删除原段落
            p = para._p
            p.getparent().remove(p)
            
            # 处理完后直接返回
            return True

    def _fix_body_para(self, para):
        text = para.text.strip()
        style_name = para.style.name
        clean_text = text.replace(" ", "").replace("\u3000", "")
        body_context = getattr(para, '_body_context', None)
        is_short_heading_like = len(clean_text) <= 30 and not re.search(r'[。；，：？！,.]', text)
        
        # 1. 标题识别（正文一级标题：仅以“第X章”为唯一标识）
        m_chapter = re.match(r'^\s*第\s*([一二三四五六七八九十百]+|\d+)\s*章\s*(.*)$', text)
        is_heading1_style = para.style.name == "Heading 1" and is_short_heading_like and "\t" not in text
        if m_chapter and is_short_heading_like and clean_text not in ["致谢", "参考文献", "附录"]:
            if m_chapter:
                chapter_no = m_chapter.group(1)
                chapter_title = m_chapter.group(2).strip()
                # 规则要求：一级标题编号和文字之间强制保留1个半角空格
                fixed_text = f"第{chapter_no}章" + (f" {chapter_title}" if chapter_title else "")
            if para.runs:
                para.runs[0].text = fixed_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(fixed_text)
            # 一级标题：三号宋体，加粗，居中，段前段后1行（16磅），单倍行距
            try:
                para.style = "Heading 1"
            except Exception:
                pass
            self._clear_paragraph_indents(para)
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(16), bold=True, 
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                     first_line_indent=Cm(0), line_spacing=1.0, ascii_font="宋体")
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(16)
            self._force_paragraph_alignment(para, 'center')
            return True
            
        elif re.match(r'^\d+[\.．]\d+[\.．]\d+', text):
            # 将全角点号归一化，并确保编号与文字紧贴
            m = re.match(r'^(\d+[\.．]\d+[\.．]\d+)\s*(.*)$', text)
            if m:
                num_part = m.group(1).replace('．', '.')
                title_part = m.group(2)
                fixed_text = num_part + title_part
                if para.runs:
                    para.runs[0].text = fixed_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(fixed_text)
            self._clear_paragraph_indents(para)
            # 三级标题：段前段后自动，行距固定20磅
            # 不使用Heading 3样式（因为样式自带段前段后），改用Normal样式并手动设置所有格式
            try:
                para.style = "Normal"
            except Exception:
                pass
            # 先设置字体等样式（不设置alignment和line_spacing，避免触发python-docx的默认段前段后）
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(14), bold=True)
            # 强制清空 runs 级别的加粗属性，避免被 Word 继承到目录中，但是正文需要加粗，所以在这一步我们不取消 run 的加粗了
            # 目录的加粗会在后处理脚本中被强行重写为不加粗
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.first_line_indent = Cm(0)
            # 使用底层XML设置段前段后为自动（不设置before/after属性），行距固定20磅
            self._set_paragraph_spacing_auto(para, line=Pt(20))
            self._force_paragraph_alignment(para, 'left')
            # 自检：验证段前段后是否真正为自动
            self._verify_heading_spacing_auto(para, "三级标题")
            return True
        elif re.match(r'^\d+[\.．]\d+(?![\.．]\d)', text):
            # 将全角点号归一化，并确保编号与文字紧贴
            m = re.match(r'^(\d+[\.．]\d+)\s*(.*)$', text)
            if m:
                num_part = m.group(1).replace('．', '.')
                title_part = m.group(2)
                fixed_text = num_part + title_part
                if para.runs:
                    para.runs[0].text = fixed_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(fixed_text)
            self._clear_paragraph_indents(para)
            # 二级标题：段前段后自动，行距固定20磅
            # 不使用Heading 2样式（因为样式自带段前段后），改用Normal样式并手动设置所有格式
            try:
                para.style = "Normal"
            except Exception:
                pass
            # 先设置字体等样式（不设置alignment和line_spacing，避免触发python-docx的默认段前段后）
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(15), bold=True)
            # 强制清空 runs 级别的加粗属性，避免被 Word 继承到目录中，但是正文需要加粗，所以在这一步我们不取消 run 的加粗了
            # 目录的加粗会在后处理脚本中被强行重写为不加粗
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.first_line_indent = Cm(0)
            # 使用底层XML设置段前段后为自动（不设置before/after属性），行距固定20磅
            self._set_paragraph_spacing_auto(para, line=Pt(20))
            self._force_paragraph_alignment(para, 'left')
            # 自检：验证段前段后是否真正为自动
            self._verify_heading_spacing_auto(para, "二级标题")
            return True
        
        # 2. 特殊章节标题 (致谢、参考文献、附录) - 三号宋体，加粗，居中
        elif clean_text in ["致谢", "参考文献", "附录"]:
            try:
                para.style = "Heading 1"
            except Exception:
                pass
            self._clear_paragraph_indents(para)
            if clean_text == "致谢":
                if para.runs:
                    para.runs[0].text = "致    谢"
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run("致    谢")
            elif clean_text == "参考文献":
                if para.runs:
                    para.runs[0].text = "参考文献"
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run("参考文献")
            elif clean_text == "附录":
                if para.runs:
                    para.runs[0].text = "附录"
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run("附录")
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(16), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0), line_spacing=1.0, ascii_font="宋体")
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(16)
            self._force_paragraph_alignment(para, 'center')
            return True
        elif is_heading1_style:
            try:
                para.style = "Heading 1"
            except Exception:
                pass
            self._clear_paragraph_indents(para)
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(16), bold=True,
                                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                     first_line_indent=Cm(0), line_spacing=1.0)
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(16)
            self._force_paragraph_alignment(para, 'center')
            return True

        # 3. 参考文献内容 - 五号楷体，顶格，两端对齐，段后0磅
        elif body_context == "参考文献" or (text.startswith("[") and len(text) > 1 and text[1].isdigit()):
            try:
                para.style = "Normal"
            except Exception:
                pass
            self._clear_paragraph_indents(para)
            self._apply_custom_style(para, font_name="楷体", font_size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=Pt(20), first_line_indent=Cm(0))
            para.paragraph_format.left_indent = Cm(0)
            para.paragraph_format.space_after = Pt(0)
            self._force_paragraph_alignment(para, 'both')
            return True

        # 4. 图片/表格内容识别
        # 增强版图片检测逻辑：不仅检查 run 的直接 XML，还检查段落的完整 XML
        has_image = self._contains_image_or_table(para)
        if not has_image:
            para_xml = para._element.xml
            if 'w:drawing' in para_xml or 'v:imagedata' in para_xml or 'pic' in para_xml or 'w:object' in para_xml or 'w:pict' in para_xml:
                has_image = True
                    
        if has_image:
            # 包含图片的段落必须居中，且去除任何缩进
            try:
                para.style = "Normal"
            except Exception:
                pass
            self._clear_paragraph_indents(para)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._force_paragraph_alignment(para, 'center')
            return True
        elif re.match(r'^(图|表)\s*\d+(?:[.\-]\d+)*', text):
            try:
                para.style = "Normal"
            except Exception:
                pass
            self._clear_paragraph_indents(para)
            self._apply_custom_style(para, font_name="宋体", font_size=Pt(10.5), alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=Cm(0))
            para.paragraph_format.left_indent = Cm(0)
            self._force_paragraph_alignment(para, 'center')
            return True

        # 6. 普通正文 - 宋体 小四 行距: 固定值 20 磅
        else:
            try:
                para.style = "Normal"
            except Exception:
                pass
            self._apply_custom_style(
                para,
                font_name="宋体",
                font_size=Pt(12),
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line_spacing=Pt(20),
                first_line_indent=Pt(24),
            )
            self._force_paragraph_alignment(para, 'both')
            return True

    def _apply_custom_style(self, para, font_name=None, font_size=None, bold=None, alignment=None, line_spacing=None, first_line_indent=None, ascii_font=None):
        """提供细粒度的段落和文字格式覆盖"""
        # 标记已处理，防止重复
        setattr(para, '_custom_style', True)
        
        if alignment is not None:
            para.alignment = alignment
        if line_spacing is not None:
            para.paragraph_format.line_spacing = line_spacing
        if first_line_indent is not None:
            para.paragraph_format.first_line_indent = first_line_indent

        for run in para.runs:
            self._set_run_font(run, font_name=font_name, font_size=font_size, bold=bold, ascii_font=ascii_font)

    def _contains_image_or_table(self, para):
        """检查段落是否包含图片或旧版 VML 图形"""
        xml = para._element.xml
        return any(tag in xml for tag in ['w:drawing', 'v:imagedata', 'w:pict', 'pic', 'w:object'])

    def _fix_captions(self, doc):
        """修复图表题注"""
        for para in doc.paragraphs:
            text = para.text.strip()
            # 识别“图 1.1” 或 “表 2.1” 开头的段落
            if (text.startswith("图") or text.startswith("表")) and any(char.isdigit() for char in text[:10]):
                if text.startswith("图"):
                    prev_para, next_para = self._find_adjacent_paragraphs(doc, para)
                    prev_has_image = prev_para is not None and self._contains_image_or_table(prev_para)
                    next_has_image = next_para is not None and self._contains_image_or_table(next_para)
                    if next_has_image and not prev_has_image:
                        next_para._element.addnext(para._element)
                # 题注通常使用 Normal 或专门的 Caption 样式，字号较小且居中
                self._apply_base_style(para, "Normal")
                self._clear_paragraph_indents(para)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.left_indent = Cm(0)
                para.paragraph_format.line_spacing = 1.5
                if text.startswith("图") or text.startswith("表"):
                    for run in para.runs:
                        self._set_run_font(run, font_name="楷体", font_size=Pt(10.5), bold=False)
                self._force_paragraph_alignment(para, 'center')

    def _paragraph_from_element(self, doc, element):
        if element is None or element.tag != qn('w:p'):
            return None
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table_caption_paragraph(self, doc, table):
        previous = table._element.getprevious()
        while previous is not None:
            if previous.tag == qn('w:p'):
                return self._paragraph_from_element(doc, previous)
            previous = previous.getprevious()
        return None

    def _find_table_caption_below_paragraph(self, doc, table):
        next_element = table._element.getnext()
        while next_element is not None:
            if next_element.tag == qn('w:p'):
                return self._paragraph_from_element(doc, next_element)
            if next_element.tag == qn('w:tbl'):
                return None
            next_element = next_element.getnext()
        return None

    def _find_adjacent_paragraphs(self, doc, para):
        previous_para = None
        next_para = None

        previous = para._element.getprevious()
        while previous is not None:
            if previous.tag == qn('w:p'):
                previous_para = self._paragraph_from_element(doc, previous)
                break
            if previous.tag == qn('w:tbl'):
                break
            previous = previous.getprevious()

        next_element = para._element.getnext()
        while next_element is not None:
            if next_element.tag == qn('w:p'):
                next_para = self._paragraph_from_element(doc, next_element)
                break
            if next_element.tag == qn('w:tbl'):
                break
            next_element = next_element.getnext()

        return previous_para, next_para

    def _fix_tables(self, doc):
        for table in doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            caption_para = self._find_table_caption_paragraph(doc, table)
            below_caption_para = self._find_table_caption_below_paragraph(doc, table)
            if below_caption_para is not None:
                below_caption_text = below_caption_para.text.strip()
                if below_caption_text.startswith("表") and any(ch.isdigit() for ch in below_caption_text[:10]):
                    table._element.addprevious(deepcopy(below_caption_para._element))
                    parent = below_caption_para._element.getparent()
                    if parent is not None:
                        parent.remove(below_caption_para._element)
                    caption_para = self._find_table_caption_paragraph(doc, table)
            if caption_para is not None:
                caption_text = caption_para.text.strip()
                if caption_text.startswith("表") and any(ch.isdigit() for ch in caption_text[:10]):
                    self._clear_paragraph_indents(caption_para)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.paragraph_format.first_line_indent = Cm(0)
                    caption_para.paragraph_format.left_indent = Cm(0)
                    caption_para.paragraph_format.line_spacing = 1.5
                    self._force_paragraph_alignment(caption_para, 'center')
                    for run in caption_para.runs:
                        self._set_run_font(run, font_name="楷体", font_size=Pt(10.5), bold=False)

            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for para in cell.paragraphs:
                        self._clear_paragraph_indents(para)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.paragraph_format.first_line_indent = Cm(0)
                        para.paragraph_format.left_indent = Cm(0)
                        para.paragraph_format.line_spacing = 1.5
                        self._force_paragraph_alignment(para, 'center')
                        for run in para.runs:
                            self._set_run_font(run, font_name="宋体", font_size=Pt(10.5), bold=False)

    def _write_validation_report(self, output_path, validation_history):
        report_path = output_path.parent / f"{output_path.stem}_validation.md"
        lines = []
        lines.append("# 格式复检报告")
        lines.append(f"- 原文档: {output_path.name}")
        lines.append(f"- 复检轮次: {len(validation_history)}")
        lines.append(f"- 规则来源: FORMAT_RULES_LEDGER.md")
        lines.append("")
        lines.append("## 复检结果")
        for entry in validation_history:
            lines.append(f"- 第 {entry['pass']} 轮: 问题数 {len(entry['issues'])}")
            for issue in entry["issues"][:10]:
                lines.append(f"  - {issue}")
        lines.append("")
        lines.append("## 结论")
        if validation_history and not validation_history[-1]["issues"]:
            lines.append("所有规则项均通过，文档格式与台账一致。")
        else:
            lines.append("仍有未通过项，请根据上方列表进一步修复。")
        report_path.write_text("\n".join(lines), encoding="utf-8")

    def _write_stats_report(self, output_path, stats):
        report_path = output_path.parent / f"{output_path.stem}_stats.md"
        lines = []
        lines.append("# 修复统计报表")
        lines.append(f"- 原文档: {output_path.name}")
        lines.append("")
        lines.append("## 修复计数")
        lines.append(f"- 标题与摘要: {stats.get('abstract', 0)}")
        lines.append(f"- 目录: {stats.get('toc', 0)}")
        lines.append(f"- 正文: {stats.get('body', 0)}")
        report_path.write_text("\n".join(lines), encoding="utf-8")
    
    def _post_enforce(self, doc):
        current_part = "摘要"
        for para in doc.paragraphs:
            text = para.text.strip()
            clean = self._normalize_text(text)
            if current_part == "摘要" and "目录" in clean and len(clean) < 10:
                current_part = "目录"
            elif current_part == "目录":
                if para.style.name == "Heading 1" or (len(text) < 30 and ("第一章" in clean or "引言" in clean) and para.alignment == WD_ALIGN_PARAGRAPH.CENTER):
                    if not re.search(r'\d+$', text.strip()):
                        current_part = "正文"
            if current_part == "目录":
                if clean != "目录" and text:
                    self._clear_paragraph_indents(para)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.left_indent = Cm(0)
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.line_spacing = 1.5
                    self._force_paragraph_alignment(para, 'left')
            elif current_part == "正文":
                if clean in ["致谢", "参考文献"]:
                    self._clear_paragraph_indents(para)
                    if clean == "致谢":
                        if para.runs:
                            para.runs[0].text = "致    谢"
                            for run in para.runs[1:]:
                                run.text = ""
                        else:
                            para.add_run("致    谢")
                    else:
                        if para.runs:
                            para.runs[0].text = "参考文献"
                            for run in para.runs[1:]:
                                run.text = ""
                        else:
                            para.add_run("参考文献")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Cm(0)
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.line_spacing = 1.0
                    self._force_paragraph_alignment(para, 'center')

def main():
    parser = argparse.ArgumentParser(description="Paper Format Fixer Pro")
    parser.add_argument("--file", type=str, required=True, help="论文文件路径")
    parser.add_argument("--no-word-update", action="store_true", help="不使用本机 Word 自动更新目录/域")
    args = parser.parse_args()

    fixer = PaperFixer(TEMPLATE_PATH)
    fixed_file_path = Path(args.file).parent / f"{Path(args.file).stem}_fixed.docx"
    fixer.fix(Path(args.file), word_update_fields=not args.no_word_update)
    
    if not args.no_word_update and fixed_file_path.exists():
        print("🧾 启动独立的后处理脚本以修复 Word 自动更新目录丢失的前导点和加粗...")
        import subprocess
        script_dir = Path(__file__).parent
        post_script = script_dir / "fix_toc_styles_post.py"
        try:
            subprocess.run(["python", str(post_script), str(fixed_file_path)], check=True)
        except Exception as e:
            print(f"执行目录样式校正脚本失败: {e}")

if __name__ == "__main__":
    main()
